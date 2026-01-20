"""块式基础计算插件UI组件"""

from PySide6.QtWidgets import (
    QWidget,
    QVBoxLayout,
    QHBoxLayout,
    QLabel,
    QLineEdit,
    QGroupBox,
    QComboBox,
    QScrollArea,
    QSizePolicy,
    QTextEdit,
    QPushButton
)
from PySide6.QtGui import QPixmap
from PySide6.QtCore import Signal, Slot, Qt

from plugins.Basic_Block.logic import BasicBlockLogic


class BasicBlockWidget(QWidget):
    """块式基础计算插件UI组件"""
    
    def __init__(self):
        """初始化UI组件"""
        super().__init__()
        
        # 初始化业务逻辑
        self._logic = BasicBlockLogic()
        
        # 初始化UI
        self._init_ui()
        
        # 连接信号和槽
        self._connect_signals()
    
    def _init_ui(self):
        """初始化UI"""
        # 创建滚动区域
        scroll_area = QScrollArea(self)
        scroll_area.setWidgetResizable(True)
        scroll_area.setVerticalScrollBarPolicy(Qt.ScrollBarAsNeeded)
        scroll_area.setHorizontalScrollBarPolicy(Qt.ScrollBarAsNeeded)
        
        # 滚动区域的容器Widget
        scroll_widget = QWidget()
        
        # 主布局：垂直布局，包含标题、参数区域和结果区域
        main_layout = QVBoxLayout(scroll_widget)
        
        # 添加标题
        title_label = QLabel("块式基础计算")
        title_label.setStyleSheet("font-size: 20px; font-weight: bold; margin-bottom: 10px;")
        title_label.setAlignment(Qt.AlignCenter)  # 设置文字居中
        main_layout.addWidget(title_label)
        
        # 参数区域：水平布局，左侧是基础尺寸+力学，右侧是换填+地脚螺栓+其它
        params_layout = QHBoxLayout()
        
        # 左侧列：基础尺寸 + 力学参数（垂直排列）
        left_column = QVBoxLayout()
        
        # 基础尺寸输入组
        basic_size_group = QGroupBox("基础尺寸")
        basic_size_group.setSizePolicy(QSizePolicy.Preferred, QSizePolicy.Fixed)
        basic_size_layout = QVBoxLayout(basic_size_group)
        
        # 基础长度
        self._length_layout = QHBoxLayout()
        self._length_layout.addWidget(QLabel("基础长度:"))
        self._length_input = QLineEdit()
        self._length_layout.addWidget(self._length_input)
        self._length_layout.addWidget(QLabel("m"))
        basic_size_layout.addLayout(self._length_layout)
        
        # 基础宽度
        self._width_layout = QHBoxLayout()
        self._width_layout.addWidget(QLabel("基础宽度:"))
        self._width_input = QLineEdit()
        self._width_layout.addWidget(self._width_input)
        self._width_layout.addWidget(QLabel("m"))
        basic_size_layout.addLayout(self._width_layout)
        
        # 基础高度
        self._height_layout = QHBoxLayout()
        self._height_layout.addWidget(QLabel("基础高度:"))
        self._height_input = QLineEdit()
        self._height_layout.addWidget(self._height_input)
        self._height_layout.addWidget(QLabel("m"))
        basic_size_layout.addLayout(self._height_layout)
        
        # 基础高出地面高度
        self._height_above_ground_layout = QHBoxLayout()
        self._height_above_ground_layout.addWidget(QLabel("基础高出地面高度:"))
        self._height_above_ground_input = QLineEdit()
        self._height_above_ground_layout.addWidget(self._height_above_ground_input)
        self._height_above_ground_layout.addWidget(QLabel("m"))
        basic_size_layout.addLayout(self._height_above_ground_layout)
        
        # 垫层厚度
        self._cushion_layout = QHBoxLayout()
        self._cushion_layout.addWidget(QLabel("垫层厚度:"))
        self._cushion_input = QLineEdit()
        self._cushion_layout.addWidget(self._cushion_input)
        self._cushion_layout.addWidget(QLabel("m"))
        basic_size_layout.addLayout(self._cushion_layout)
        
        # 二次灌浆厚度
        self._grout_layout = QHBoxLayout()
        self._grout_layout.addWidget(QLabel("二次灌浆厚度:"))
        self._grout_input = QLineEdit()
        self._grout_layout.addWidget(self._grout_input)
        self._grout_layout.addWidget(QLabel("mm"))
        basic_size_layout.addLayout(self._grout_layout)
        
        # 基础混凝土材质选项
        self._concrete_material_layout = QHBoxLayout()
        self._concrete_material_layout.addWidget(QLabel("基础混凝土材质:"))
        self._concrete_material_combo = QComboBox()
        self._concrete_material_combo.addItems(["C20", "C25", "C30", "C35", "C40"])
        self._concrete_material_layout.addWidget(self._concrete_material_combo)
        basic_size_layout.addLayout(self._concrete_material_layout)
        
        left_column.addWidget(basic_size_group)
        
        # 力学参数输入组
        mechanics_group = QGroupBox("力学参数")
        mechanics_group.setSizePolicy(QSizePolicy.Preferred, QSizePolicy.Fixed)
        mechanics_layout = QVBoxLayout(mechanics_group)
        
        # 地基承载力
        self._bearing_layout = QHBoxLayout()
        self._bearing_layout.addWidget(QLabel("地基承载力:"))
        self._bearing_input = QLineEdit()
        self._bearing_layout.addWidget(self._bearing_input)
        self._bearing_layout.addWidget(QLabel("kPa"))
        mechanics_layout.addLayout(self._bearing_layout)
        
        # 上部荷载
        self._load_layout = QHBoxLayout()
        self._load_layout.addWidget(QLabel("上部荷载:"))
        self._load_input = QLineEdit()
        self._load_layout.addWidget(self._load_input)
        self._load_layout.addWidget(QLabel("KN"))
        mechanics_layout.addLayout(self._load_layout)
        
        left_column.addWidget(mechanics_group)
        
        # 预埋钢板参数输入组
        embedded_plate_group = QGroupBox("预埋钢板参数")
        embedded_plate_group.setSizePolicy(QSizePolicy.Preferred, QSizePolicy.Fixed)
        embedded_plate_layout = QHBoxLayout(embedded_plate_group)
        
        # 预埋钢板长度
        self._plate_length_layout = QHBoxLayout()
        self._plate_length_layout.addWidget(QLabel("长度:"))
        self._plate_length_input = QLineEdit()
        self._plate_length_layout.addWidget(self._plate_length_input)
        self._plate_length_layout.addWidget(QLabel("m"))
        embedded_plate_layout.addLayout(self._plate_length_layout)
        
        # 预埋钢板宽度
        self._plate_width_layout = QHBoxLayout()
        self._plate_width_layout.addWidget(QLabel("宽度:"))
        self._plate_width_input = QLineEdit()
        self._plate_width_layout.addWidget(self._plate_width_input)
        self._plate_width_layout.addWidget(QLabel("m"))
        embedded_plate_layout.addLayout(self._plate_width_layout)
        
        # 预埋钢板厚度
        self._plate_thickness_layout = QHBoxLayout()
        self._plate_thickness_layout.addWidget(QLabel("厚度:"))
        self._plate_thickness_input = QLineEdit()
        self._plate_thickness_layout.addWidget(self._plate_thickness_input)
        self._plate_thickness_layout.addWidget(QLabel("mm"))
        embedded_plate_layout.addLayout(self._plate_thickness_layout)
        
        left_column.addWidget(embedded_plate_group)
        
        # 添加伸缩空间，使左侧列高度与右侧列保持一致
        left_column.addStretch()
        
        # 右侧列：换填级配砂石参数 + 地脚螺栓参数 + 其它参数（垂直排列）
        right_column = QVBoxLayout()
        
        # 换填级配砂石参数输入组
        replacement_group = QGroupBox("换填级配砂石参数")
        replacement_group.setSizePolicy(QSizePolicy.Preferred, QSizePolicy.Fixed)
        replacement_layout = QVBoxLayout(replacement_group)
        
        # 换填厚度
        self._replacement_thickness_layout = QHBoxLayout()
        self._replacement_thickness_layout.addWidget(QLabel("换填厚度:"))
        self._replacement_thickness_input = QLineEdit()
        self._replacement_thickness_layout.addWidget(self._replacement_thickness_input)
        self._replacement_thickness_layout.addWidget(QLabel("m"))
        replacement_layout.addLayout(self._replacement_thickness_layout)
        
        # 换填宽度
        self._replacement_width_layout = QHBoxLayout()
        self._replacement_width_layout.addWidget(QLabel("换填宽度:"))
        self._replacement_width_input = QLineEdit()
        self._replacement_width_layout.addWidget(self._replacement_width_input)
        self._replacement_width_layout.addWidget(QLabel("m"))
        replacement_layout.addLayout(self._replacement_width_layout)
        
        right_column.addWidget(replacement_group)
        
        # 地脚螺栓参数输入组
        anchor_group = QGroupBox("地脚螺栓参数")
        anchor_group.setSizePolicy(QSizePolicy.Preferred, QSizePolicy.Fixed)
        anchor_layout = QVBoxLayout(anchor_group)
        
        # 地脚螺栓个数
        self._anchor_count_layout = QHBoxLayout()
        self._anchor_count_layout.addWidget(QLabel("地脚螺栓个数:"))
        self._anchor_count_input = QLineEdit()
        self._anchor_count_layout.addWidget(self._anchor_count_input)
        self._anchor_count_layout.addWidget(QLabel("个"))
        anchor_layout.addLayout(self._anchor_count_layout)
        
        # 地脚螺栓直径
        self._anchor_diam_layout = QHBoxLayout()
        self._anchor_diam_layout.addWidget(QLabel("地脚螺栓直径:"))
        self._anchor_diam_input = QLineEdit()
        self._anchor_diam_layout.addWidget(self._anchor_diam_input)
        self._anchor_diam_layout.addWidget(QLabel("mm"))
        anchor_layout.addLayout(self._anchor_diam_layout)
        
        # 地脚螺栓长度
        self._anchor_length_layout = QHBoxLayout()
        self._anchor_length_layout.addWidget(QLabel("地脚螺栓长度:"))
        self._anchor_length_input = QLineEdit()
        self._anchor_length_layout.addWidget(self._anchor_length_input)
        self._anchor_length_layout.addWidget(QLabel("mm"))
        anchor_layout.addLayout(self._anchor_length_layout)
        

        
        right_column.addWidget(anchor_group)
        
        # 其它参数输入组
        other_group = QGroupBox("其它参数")
        other_group.setSizePolicy(QSizePolicy.Preferred, QSizePolicy.Fixed)
        other_layout = QVBoxLayout(other_group)
        
        # 是否素砼
        self._is_plain_concrete_layout = QHBoxLayout()
        self._is_plain_concrete_layout.addWidget(QLabel("是否素砼?"))
        
        # 创建按钮容器
        button_container = QWidget()
        button_layout = QHBoxLayout(button_container)
        button_layout.setSpacing(10)
        button_layout.setContentsMargins(0, 0, 0, 0)
        
        # 否按钮
        self._is_plain_concrete_no_btn = QPushButton("否")
        self._is_plain_concrete_no_btn.setCheckable(True)
        self._is_plain_concrete_no_btn.setChecked(True)  # 默认否
        self._is_plain_concrete_no_btn.setFixedWidth(80)
        self._is_plain_concrete_no_btn.setStyleSheet("""
            QPushButton {
                background: qlineargradient(x1:0, y1:0, x2:0, y2:1, stop:0 #2c2c2c, stop:1 #1e1e1e);
                color: #e0e0e0;
                border: 1px solid #444;
                border-radius: 15px;
                padding: 8px 15px;
                text-align: center;
                font-size: 14px;
            }
            QPushButton:hover {
                background: qlineargradient(x1:0, y1:0, x2:0, y2:1, stop:0 #40a1e0, stop:1 #3498DB);
                border: 1px solid #2980B9;
            }
            QPushButton:checked {
                background: qlineargradient(x1:0, y1:0, x2:0, y2:1, stop:0 #3498DB, stop:1 #2980B9);
                color: white;
                border: 1px solid #2471A3;
            }
        """)
        button_layout.addWidget(self._is_plain_concrete_no_btn)
        
        # 是按钮
        self._is_plain_concrete_yes_btn = QPushButton("是")
        self._is_plain_concrete_yes_btn.setCheckable(True)
        self._is_plain_concrete_yes_btn.setFixedWidth(80)
        self._is_plain_concrete_yes_btn.setStyleSheet("""
            QPushButton {
                background: qlineargradient(x1:0, y1:0, x2:0, y2:1, stop:0 #2c2c2c, stop:1 #1e1e1e);
                color: #e0e0e0;
                border: 1px solid #444;
                border-radius: 15px;
                padding: 8px 15px;
                text-align: center;
                font-size: 14px;
            }
            QPushButton:hover {
                background: qlineargradient(x1:0, y1:0, x2:0, y2:1, stop:0 #40a1e0, stop:1 #3498DB);
                border: 1px solid #2980B9;
            }
            QPushButton:checked {
                background: qlineargradient(x1:0, y1:0, x2:0, y2:1, stop:0 #3498DB, stop:1 #2980B9);
                color: white;
                border: 1px solid #2471A3;
            }
        """)
        button_layout.addWidget(self._is_plain_concrete_yes_btn)
        
        # 添加到布局
        self._is_plain_concrete_layout.addWidget(button_container)
        other_layout.addLayout(self._is_plain_concrete_layout)
        
        # 是否打桩
        self._is_pile_layout = QHBoxLayout()
        self._is_pile_layout.addWidget(QLabel("是否打桩?"))
        
        # 创建按钮容器
        pile_button_container = QWidget()
        pile_button_layout = QHBoxLayout(pile_button_container)
        pile_button_layout.setSpacing(10)
        pile_button_layout.setContentsMargins(0, 0, 0, 0)
        
        # 否按钮
        self._is_pile_no_btn = QPushButton("否")
        self._is_pile_no_btn.setCheckable(True)
        self._is_pile_no_btn.setChecked(True)  # 默认否
        self._is_pile_no_btn.setFixedWidth(80)
        self._is_pile_no_btn.setStyleSheet("""
            QPushButton {
                background: qlineargradient(x1:0, y1:0, x2:0, y2:1, stop:0 #2c2c2c, stop:1 #1e1e1e);
                color: #e0e0e0;
                border: 1px solid #444;
                border-radius: 15px;
                padding: 8px 15px;
                text-align: center;
                font-size: 14px;
            }
            QPushButton:hover {
                background: qlineargradient(x1:0, y1:0, x2:0, y2:1, stop:0 #40a1e0, stop:1 #3498DB);
                border: 1px solid #2980B9;
            }
            QPushButton:checked {
                background: qlineargradient(x1:0, y1:0, x2:0, y2:1, stop:0 #3498DB, stop:1 #2980B9);
                color: white;
                border: 1px solid #2471A3;
            }
        """)
        pile_button_layout.addWidget(self._is_pile_no_btn)
        
        # 是按钮
        self._is_pile_yes_btn = QPushButton("是")
        self._is_pile_yes_btn.setCheckable(True)
        self._is_pile_yes_btn.setFixedWidth(80)
        self._is_pile_yes_btn.setStyleSheet("""
            QPushButton {
                background: qlineargradient(x1:0, y1:0, x2:0, y2:1, stop:0 #2c2c2c, stop:1 #1e1e1e);
                color: #e0e0e0;
                border: 1px solid #444;
                border-radius: 15px;
                padding: 8px 15px;
                text-align: center;
                font-size: 14px;
            }
            QPushButton:hover {
                background: qlineargradient(x1:0, y1:0, x2:0, y2:1, stop:0 #40a1e0, stop:1 #3498DB);
                border: 1px solid #2980B9;
            }
            QPushButton:checked {
                background: qlineargradient(x1:0, y1:0, x2:0, y2:1, stop:0 #3498DB, stop:1 #2980B9);
                color: white;
                border: 1px solid #2471A3;
            }
        """)
        pile_button_layout.addWidget(self._is_pile_yes_btn)
        
        # 添加到布局
        self._is_pile_layout.addWidget(pile_button_container)
        other_layout.addLayout(self._is_pile_layout)
        
        # 桩根数
        self._pile_count_layout = QHBoxLayout()
        self._pile_count_layout.addWidget(QLabel("桩根数:"))
        self._pile_count_input = QLineEdit()
        self._pile_count_input.setEnabled(False)  # 默认禁用
        self._pile_count_layout.addWidget(self._pile_count_input)
        self._pile_count_layout.addWidget(QLabel("根"))
        other_layout.addLayout(self._pile_count_layout)
        
        # 基础个数
        self._foundation_count_layout = QHBoxLayout()
        self._foundation_count_layout.addWidget(QLabel("基础个数:"))
        self._foundation_count_input = QLineEdit()
        self._foundation_count_layout.addWidget(self._foundation_count_input)
        self._foundation_count_layout.addWidget(QLabel("个"))
        other_layout.addLayout(self._foundation_count_layout)
        
        right_column.addWidget(other_group)
        
        # 将左右列添加到参数布局
        params_layout.addLayout(left_column)
        params_layout.addLayout(right_column)
        
        # 将参数布局添加到主布局
        main_layout.addLayout(params_layout)
        
        # 计算按钮区域
        calculate_layout = QHBoxLayout()
        
        # 计算按钮
        self._calculate_btn = QPushButton("计算")
        self._calculate_btn.setStyleSheet("""
            QPushButton {
                background: qlineargradient(x1:0, y1:0, x2:0, y2:1, stop:0 #3498DB, stop:1 #2980B9);
                color: white;
                border: 1px solid #2471A3;
                border-radius: 8px;
                padding: 10px 20px;
                font-size: 14px;
                min-width: 100px;
            }
            QPushButton:hover {
                background: qlineargradient(x1:0, y1:0, x2:0, y2:1, stop:0 #3aade0, stop:1 #3498DB);
                border: 1px solid #2980B9;
            }
            QPushButton:pressed {
                background: qlineargradient(x1:0, y1:0, x2:0, y2:1, stop:0 #2471A3, stop:1 #1f638c);
            }
        """)
        calculate_layout.addWidget(self._calculate_btn)
        
        # 导出计算书按钮
        self._export_btn = QPushButton("导出计算书")
        self._export_btn.setStyleSheet("""
            QPushButton {
                background: qlineargradient(x1:0, y1:0, x2:0, y2:1, stop:0 #2ecc71, stop:1 #27ae60);
                color: white;
                border: 1px solid #229954;
                border-radius: 8px;
                padding: 10px 20px;
                font-size: 14px;
                min-width: 120px;
                margin-left: 10px;
            }
            QPushButton:hover {
                background: qlineargradient(x1:0, y1:0, x2:0, y2:1, stop:0 #32e17c, stop:1 #2ecc71);
                border: 1px solid #27ae60;
            }
            QPushButton:pressed {
                background: qlineargradient(x1:0, y1:0, x2:0, y2:1, stop:0 #229954, stop:1 #1e8449);
            }
        """)
        calculate_layout.addWidget(self._export_btn)

        # 导出料表按钮
        self._export_material_btn = QPushButton("导出料表")
        self._export_material_btn.setStyleSheet("""
            QPushButton {
                background: qlineargradient(x1:0, y1:0, x2:0, y2:1, stop:0 #3498db, stop:1 #2980b9);
                color: white;
                border: 1px solid #2471A3;
                border-radius: 8px;
                padding: 10px 20px;
                font-size: 14px;
                min-width: 120px;
                margin-left: 10px;
            }
            QPushButton:hover {
                background: qlineargradient(x1:0, y1:0, x2:0, y2:1, stop:0 #5dade2, stop:1 #3498db);
                border: 1px solid #2980b9;
            }
            QPushButton:pressed {
                background: qlineargradient(x1:0, y1:0, x2:0, y2:1, stop:0 #2471A3, stop:1 #1f638b);
            }
        """)
        calculate_layout.addWidget(self._export_material_btn)
        calculate_layout.addStretch()  # 添加弹簧使按钮左对齐
        main_layout.addLayout(calculate_layout)
        
        # 输出结果显示区域：放在主布局的最下面，通过滚动条访问
        result_group = QGroupBox("输出结果")
        result_layout = QVBoxLayout(result_group)
        
        # 结果显示文本框
        self._result_text = QTextEdit()
        self._result_text.setReadOnly(True)
        self._result_text.setMinimumHeight(400)  # 输出结果显示框的最小高度
        # 设置自定义上下文菜单策略
        self._result_text.setContextMenuPolicy(Qt.ContextMenuPolicy.CustomContextMenu)
        # 连接上下文菜单信号
        self._result_text.customContextMenuRequested.connect(self._on_result_text_context_menu)
        result_layout.addWidget(self._result_text)
        
        # 将结果区域添加到主布局的最下面
        main_layout.addWidget(result_group)
        
        # 设置滚动区域的Widget
        scroll_area.setWidget(scroll_widget)
        
        # 设置主布局
        main_widget_layout = QVBoxLayout(self)
        main_widget_layout.addWidget(scroll_area)
        
        # 为所有QLineEdit控件设置自定义上下文菜单
        for widget in self.findChildren(QLineEdit):
            widget.setContextMenuPolicy(Qt.ContextMenuPolicy.CustomContextMenu)
            widget.customContextMenuRequested.connect(self._on_input_context_menu)
        
    def _connect_signals(self):
        """连接信号和槽"""
        # 连接是否素砼按钮的互斥信号
        self._is_plain_concrete_no_btn.clicked.connect(self._on_plain_concrete_button_clicked)
        self._is_plain_concrete_yes_btn.clicked.connect(self._on_plain_concrete_button_clicked)
        # 连接是否打桩按钮的互斥信号
        self._is_pile_no_btn.clicked.connect(self._on_pile_button_clicked)
        self._is_pile_yes_btn.clicked.connect(self._on_pile_button_clicked)
        # 连接计算按钮的点击信号
        self._calculate_btn.clicked.connect(self._on_calculate_clicked)
        # 连接导出计算书按钮的点击信号
        self._export_btn.clicked.connect(self._on_export_triggered)
        # 连接导出料表按钮的点击信号
        self._export_material_btn.clicked.connect(self._on_export_material_triggered)
    
    @Slot()
    def _on_result_text_context_menu(self, pos):
        """处理输出结果文本框的上下文菜单请求"""
        from PySide6.QtWidgets import QMenu
        
        # 创建上下文菜单
        menu = QMenu(self._result_text)
        
        # 添加复制选项
        copy_action = menu.addAction("复制 (Ctrl+C)")
        copy_action.triggered.connect(self._result_text.copy)
        
        # 添加全选选项
        select_all_action = menu.addAction("全选 (Ctrl+A)")
        select_all_action.triggered.connect(self._result_text.selectAll)
        
        # 显示菜单
        menu.exec_(self._result_text.mapToGlobal(pos))
    
    @Slot()
    def _on_input_context_menu(self, pos):
        """处理输入控件的上下文菜单请求"""
        from PySide6.QtWidgets import QMenu
        
        # 获取发送信号的控件
        sender = self.sender()
        if not sender:
            return
        
        # 创建上下文菜单
        menu = QMenu(sender)
        
        # 添加撤销选项
        undo_action = menu.addAction("撤销 (Ctrl+Z)")
        undo_action.triggered.connect(sender.undo)
        
        # 添加重做选项
        redo_action = menu.addAction("重做 (Ctrl+Y)")
        redo_action.triggered.connect(sender.redo)
        
        menu.addSeparator()
        
        # 添加剪切选项
        cut_action = menu.addAction("剪切 (Ctrl+X)")
        cut_action.triggered.connect(sender.cut)
        
        # 添加复制选项
        copy_action = menu.addAction("复制 (Ctrl+C)")
        copy_action.triggered.connect(sender.copy)
        
        # 添加粘贴选项
        paste_action = menu.addAction("粘贴 (Ctrl+V)")
        paste_action.triggered.connect(sender.paste)
        
        menu.addSeparator()
        
        # 添加删除选项
        delete_action = menu.addAction("删除")
        delete_action.triggered.connect(sender.del_)
        
        menu.addSeparator()
        
        # 添加全选选项
        select_all_action = menu.addAction("全选 (Ctrl+A)")
        select_all_action.triggered.connect(sender.selectAll)
        
        # 显示菜单
        menu.exec_(sender.mapToGlobal(pos))
    
    @Slot()
    def _on_plain_concrete_button_clicked(self):
        """是否素砼按钮点击事件，确保互斥"""
        sender = self.sender()
        if sender == self._is_plain_concrete_no_btn:
            self._is_plain_concrete_yes_btn.setChecked(False)
        else:
            self._is_plain_concrete_no_btn.setChecked(False)
    
    @Slot()
    def _on_pile_button_clicked(self):
        """是否打桩按钮点击事件，确保互斥并控制桩根数输入框的启用/禁用"""
        sender = self.sender()
        if sender == self._is_pile_no_btn:
            self._is_pile_yes_btn.setChecked(False)
            self._pile_count_input.setEnabled(False)  # 禁用桩根数输入框
            self._pile_count_input.clear()  # 清空桩根数输入框
        else:
            self._is_pile_no_btn.setChecked(False)
            self._pile_count_input.setEnabled(True)  # 启用桩根数输入框
    
    @Slot()
    def _on_calculate_clicked(self):
        """计算按钮点击事件"""
        try:
            # 获取输入值并转换为数值类型
            length = float(self._length_input.text()) if self._length_input.text() else 0
            width = float(self._width_input.text()) if self._width_input.text() else 0
            height = float(self._height_input.text()) if self._height_input.text() else 0
            height_above_ground = float(self._height_above_ground_input.text()) if self._height_above_ground_input.text() else 0
            # 计算基底埋深：基础高度减去基础高出地面高度
            depth = height - height_above_ground
            cushion_thickness = float(self._cushion_input.text()) if self._cushion_input.text() else 0
            grout_thickness = float(self._grout_input.text()) if self._grout_input.text() else 0
            replacement_width = float(self._replacement_width_input.text()) if self._replacement_width_input.text() else 0
            replacement_thickness = float(self._replacement_thickness_input.text()) if self._replacement_thickness_input.text() else 0
            pile_count = int(self._pile_count_input.text()) if self._pile_count_input.text() else 0
            
            # 获取基础个数
            foundation_count = int(self._foundation_count_input.text()) if self._foundation_count_input.text() else 1  # 默认1个
            
            # 获取地脚螺栓相关参数
            anchor_count = int(self._anchor_count_input.text()) if self._anchor_count_input.text() else 0
            anchor_diam = float(self._anchor_diam_input.text()) if self._anchor_diam_input.text() else 0
            anchor_length = float(self._anchor_length_input.text()) if self._anchor_length_input.text() else 0
            
            # 调用计算方法（单个基础）
            basic_volume_single = self._logic.calculate_basic_volume(length, width, height)
            cushion_volume_single = self._logic.calculate_cushion_volume(length, width, cushion_thickness)
            replacement_volume_single = self._logic.calculate_replacement_volume(length, width, replacement_width, replacement_thickness)
            
            # 获取预埋钢板参数
            plate_length = float(self._plate_length_input.text()) if self._plate_length_input.text() else 0
            plate_width = float(self._plate_width_input.text()) if self._plate_width_input.text() else 0
            plate_thickness = float(self._plate_thickness_input.text()) if self._plate_thickness_input.text() else 0
            
            # 计算预埋钢板体积（单个基础）
            plate_volume_single = self._logic.calculate_plate_volume(plate_length, plate_width, plate_thickness)
            
            # 计算二次灌浆体积（单个基础）
            grout_volume_single = self._logic.calculate_grout_volume(length, width, grout_thickness)
            
            # 计算钢材重量（单个基础）
            anchor_bolt_volume = self._logic.calculate_anchor_bolt_volume(anchor_diam, anchor_length)
            # 总钢材体积 = 地脚螺栓体积 + 预埋钢板体积
            total_steel_volume_single = anchor_bolt_volume * anchor_count + plate_volume_single
            steel_weight_kg_single = total_steel_volume_single * 7850  # 钢材密度：7850 kg/m³
            steel_weight_single = steel_weight_kg_single / 1000  # 转换为吨(t)
            
            # 乘以基础个数，得到最终结果
            basic_volume = basic_volume_single * foundation_count
            cushion_volume = cushion_volume_single * foundation_count
            replacement_volume = replacement_volume_single * foundation_count
            steel_weight = steel_weight_single * foundation_count
            steel_weight_kg = steel_weight_kg_single * foundation_count
            grout_volume = grout_volume_single * foundation_count
            plate_volume = plate_volume_single * foundation_count
            
            # 获取是否素砼
            is_plain_concrete = self._is_plain_concrete_yes_btn.isChecked()
            
            # 验算地基承载力（单个基础）
            bearing_check_result_single = self._logic.check_bearing_capacity(
                float(self._load_input.text()) if self._load_input.text() else 0,
                length,
                width,
                height,
                float(self._bearing_input.text()) if self._bearing_input.text() else 0,
                is_plain_concrete
            )
            
            # 计算总基础的地基承载力验算（对于单个基础，结果相同）
            is_bearing_satisfied, basic_weight, total_load, base_pressure, concrete_density = bearing_check_result_single
            
            # 计算基础防腐面积（单个基础）
            anticorrosion_area_single = self._logic.calculate_anticorrosion_area(length, width, depth)
            # 乘以基础个数，得到最终结果
            anticorrosion_area = anticorrosion_area_single * foundation_count
            
            # 格式化输出结果为HTML格式
            result_html = f"""<html>
            <head>
                <style>
                    body {{ font-family: 'Microsoft YaHei UI', 'Consolas', 'Segoe UI', sans-serif; font-size: 14px; line-height: 1.6; }}
                    
                    /* 通用卡片样式 */
                    .card {{ background: qlineargradient(x1:0, y1:0, x2:0, y2:1, stop:0 #ffffff, stop:1 #f5f5f5); border: 1px solid #ccc; border-radius: 8px; padding: 15px; margin-bottom: 15px; }}
                    .dark-theme .card {{ background: qlineargradient(x1:0, y1:0, x2:0, y2:1, stop:0 #2c2c2c, stop:1 #1e1e1e); border: 1px solid #444; }}
                    
                    /* 标题样式 */
                    h2 {{ color: #3498DB; margin-top: 0; margin-bottom: 15px; font-size: 16px; font-weight: bold; border-bottom: 2px solid #3498DB; padding-bottom: 5px; }}
                    .dark-theme h2 {{ color: #3498DB; border-bottom-color: #3498DB; }}
                    
                    /* 参数表格样式 */
                    .param-table {{ width: 100%; border-collapse: collapse; }}
                    .param-table tr {{ border-bottom: 1px solid #eee; }}
                    .dark-theme .param-table tr {{ border-bottom-color: #333; }}
                    .param-table td {{ padding: 8px 15px; vertical-align: middle; }}
                    
                    /* 参数项样式 */
                    .param-item {{ width: 33.33%; padding: 8px 15px; }}
                    
                    /* 标签和值样式 */
                    .param-label {{ font-weight: bold; color: #555; margin-right: 8px; }}
                    .dark-theme .param-label {{ color: #ffffff; }}
                    .param-value {{ color: #333; }}
                    .dark-theme .param-value {{ color: #ffffff; }}
                    
                    /* 计算结果区样式 */
                    .result-section {{ background: qlineargradient(x1:0, y1:0, x2:0, y2:1, stop:0 #e8f4f8, stop:1 #d4e6f1); border: 1px solid #aed6f1; border-radius: 8px; padding: 20px; margin-bottom: 15px; }}
                    .dark-theme .result-section {{ background: qlineargradient(x1:0, y1:0, x2:0, y2:1, stop:0 #1a252f, stop:1 #0f171d); border: 1px solid #3498DB; }}
                    
                    .calculation {{ margin-bottom: 15px; padding: 10px 0; background: transparent; border-radius: 0; }}
                    .dark-theme .calculation {{ background: transparent; }}
                    
                    .formula {{ font-family: 'Consolas', monospace; color: #2c3e50; margin-bottom: 5px; }}
                    .dark-theme .formula {{ color: #ffffff; }}
                    
                    .formula-result {{ font-weight: bold; color: #3498DB; margin-left: 20px; }}
                    .dark-theme .formula-result {{ color: #64b5f6; }}
                    
                    /* 最终结果样式 */
                    .final-results {{ background: qlineargradient(x1:0, y1:0, x2:0, y2:1, stop:0 #f8f9fa, stop:1 #e9ecef); border: 1px solid #dee2e6; border-radius: 8px; padding: 15px; }}
                    .dark-theme .final-results {{ background: qlineargradient(x1:0, y1:0, x2:0, y2:1, stop:0 #2c3e50, stop:1 #1a2530); border: 1px solid #444; }}
                    
                    .final-table {{ width: 100%; border-collapse: collapse; }}
                    .final-table tr {{ border-bottom: 1px solid #eee; }}
                    .dark-theme .final-table tr {{ border-bottom-color: #333; }}
                    .final-table td {{ padding: 10px 15px; vertical-align: middle; }}
                    .result-label {{ font-weight: bold; color: #333; margin-right: 15px; min-width: 180px; }}
                    .dark-theme .result-label {{ color: #ffffff; }}
                    .result-value {{ font-weight: bold; font-size: 16px; color: #3498DB; }}
                    .dark-theme .result-value {{ color: #64b5f6; }}
                </style>
            </head>
            <body>
                <!-- 最终结果区 -->
                <div class="final-results">
                    <h2>最终计算结果</h2>
                    <table class="final-table">
                        <tr><td class="result-label">基础体积</td><td class="result-value">{basic_volume:.4f} m³</td></tr>
                        <tr><td class="result-label">垫层体积</td><td class="result-value">{cushion_volume:.4f} m³</td></tr>
                        <tr><td class="result-label">换填级配砂石体积</td><td class="result-value">{replacement_volume:.4f} m³</td></tr>
                        <tr><td class="result-label">二次灌浆体积</td><td class="result-value">{grout_volume:.4f} m³</td></tr>
                        <tr><td class="result-label">钢材重量</td><td class="result-value">{steel_weight:.4f} t</td></tr>
                        <tr><td class="result-label">基础防腐面积</td><td class="result-value">{anticorrosion_area:.4f} m²</td></tr>
                        <tr><td class="result-label">地基承载力验算</td><td class="result-value">{'✅ 满足要求' if is_bearing_satisfied else '❌ 不满足要求'}</td></tr>
                    </table>
                </div>
                
                <!-- 输入参数区 -->
                <div class="card">
                    <h2>输入参数</h2>
                    <table class="param-table">
                        <tr>
                            <td class="param-item"><span class="param-label">基础长度：</span><span class="param-value">{length}m</span></td>
                            <td class="param-item"><span class="param-label">基础宽度：</span><span class="param-value">{width}m</span></td>
                            <td class="param-item"><span class="param-label">基础高度：</span><span class="param-value">{height}m</span></td>
                        </tr>
                        <tr>
                            <td class="param-item"><span class="param-label">基础高出地面高度：</span><span class="param-value">{height_above_ground}m</span></td>
                            <td class="param-item"><span class="param-label">基底埋深：</span><span class="param-value">{depth:.4f}m</span></td>
                            <td class="param-item"><span class="param-label">二次灌浆厚度：</span><span class="param-value">{grout_thickness}mm</span></td>
                        </tr>
                        <tr>
                            <td class="param-item"><span class="param-label">地基承载力：</span><span class="param-value">{self._bearing_input.text()}kPa</span></td>
                            <td class="param-item"><span class="param-label">上部荷载：</span><span class="param-value">{self._load_input.text()}KN</span></td>
                            <td class="param-item"><span class="param-label">预埋钢板：</span><span class="param-value">{plate_length}m × {plate_width}m × {plate_thickness}mm</span></td>
                        </tr>
                        <tr>
                            <td class="param-item"><span class="param-label">地脚螺栓个数：</span><span class="param-value">{self._anchor_count_input.text()}个</span></td>
                            <td class="param-item"><span class="param-label">地脚螺栓直径：</span><span class="param-value">{self._anchor_diam_input.text()}mm</span></td>
                            <td class="param-item"><span class="param-label">地脚螺栓长度：</span><span class="param-value">{self._anchor_length_input.text()}mm</span></td>
                        </tr>
                        <tr>
                            <td class="param-item"><span class="param-label">基础混凝土材质：</span><span class="param-value">{self._concrete_material_combo.currentText()}</span></td>
                            <td class="param-item"><span class="param-label">是否素砼：</span><span class="param-value">{'是' if self._is_plain_concrete_yes_btn.isChecked() else '否'}</span></td>
                            <td class="param-item"></td>
                        </tr>
                        <tr>
                            <td class="param-item"><span class="param-label">是否打桩：</span><span class="param-value">{'是' if self._is_pile_yes_btn.isChecked() else '否'}</span></td>
                            <td class="param-item"><span class="param-label">桩根数：</span><span class="param-value">{self._pile_count_input.text()}根</span></td>
                            <td class="param-item"></td>
                        </tr>
                        <tr>
                            <td class="param-item"><span class="param-label">换填厚度：</span><span class="param-value">{replacement_thickness}m</span></td>
                            <td class="param-item"><span class="param-label">换填宽度：</span><span class="param-value">{replacement_width}m</span></td>
                            <td class="param-item"></td>
                        </tr>
                    </table>
                </div>
                
                <!-- 计算结果区 -->
                <div class="result-section">
                    <h2>计算过程</h2>
                    
                    <div class="calculation">
                        <div class="formula">一、基础体积计算</div>
                        <div class="formula">（一）单个基础体积：{length}m × {width}m × {height}m</div>
                        <div class="formula-result">= {basic_volume_single:.4f} m³</div>
                        <div class="formula">（二）总基础体积：{basic_volume_single:.4f}m³ × {foundation_count}个</div>
                        <div class="formula-result">= {basic_volume:.4f} m³</div>
                    </div>
                    
                    <div class="calculation">
                        <div class="formula">二、垫层体积计算</div>
                        <div class="formula">（一）单个基础垫层体积：({length}m + 2×0.1m) × ({width}m + 2×0.1m) × {cushion_thickness}m</div>
                        <div class="formula-result">= {cushion_volume_single:.4f} m³</div>
                        <div class="formula">（二）总垫层体积：{cushion_volume_single:.4f}m³ × {foundation_count}个</div>
                        <div class="formula-result">= {cushion_volume:.4f} m³</div>
                    </div>
                    
                    <div class="calculation">
                        <div class="formula">三、换填级配砂石体积计算</div>
                        <div class="formula">（一）单个基础换填级配砂石体积：({length}m + 2 × {replacement_width}m) × ({width}m + 2 × {replacement_width}m) × {replacement_thickness}m</div>
                        <div class="formula-result">= {replacement_volume_single:.4f} m³</div>
                        <div class="formula">（二）总换填级配砂石体积：{replacement_volume_single:.4f}m³ × {foundation_count}个</div>
                        <div class="formula-result">= {replacement_volume:.4f} m³</div>
                    </div>
                    
                    <div class="calculation">
                        <div class="formula">四、二次灌浆体积计算</div>
                        <div class="formula">（一）单个基础二次灌浆体积：{length}m × {width}m × {grout_thickness}mm</div>
                        <div class="formula">= {length}m × {width}m × {grout_thickness/1000:.6f}m</div>
                        <div class="formula-result">= {grout_volume_single:.4f} m³</div>
                        <div class="formula">（二）总二次灌浆体积：{grout_volume_single:.4f}m³ × {foundation_count}个</div>
                        <div class="formula-result">= {grout_volume:.4f} m³</div>
                    </div>
                    
                    <div class="calculation">
                        <div class="formula">五、预埋钢板体积计算</div>
                        <div class="formula">（一）单个预埋钢板体积：{plate_length}m × {plate_width}m × {plate_thickness}mm</div>
                        <div class="formula">= {plate_length}m × {plate_width}m × {plate_thickness/1000:.6f}m</div>
                        <div class="formula-result">= {plate_volume_single:.6f} m³</div>
                        <div class="formula">（二）总预埋钢板体积：{plate_volume_single:.6f}m³ × {foundation_count}个</div>
                        <div class="formula-result">= {plate_volume:.4f} m³</div>
                    </div>
                    
                    <div class="calculation">
                        <div class="formula">六、钢材重量计算</div>
                        <div class="formula">（一）单根地脚螺栓体积：π × ({anchor_diam/1000/2:.4f}m)² × {anchor_length/1000:.4f}m</div>
                        <div class="formula-result">= {anchor_bolt_volume:.8f} m³</div>
                        <div class="formula">（二）单根地脚螺栓总重量：{anchor_bolt_volume:.8f}m³ × {anchor_count}根</div>
                        <div class="formula-result">= {anchor_bolt_volume * anchor_count:.8f} m³</div>
                        <div class="formula">（三）预埋钢板体积：{plate_length}m × {plate_width}m × {plate_thickness/1000:.6f}m</div>
                        <div class="formula-result">= {plate_volume_single:.6f} m³</div>
                        <div class="formula">（四）单个基础总钢材体积：地脚螺栓体积 + 预埋钢板体积</div>
                        <div class="formula">= {anchor_bolt_volume * anchor_count:.8f}m³ + {plate_volume_single:.6f}m³</div>
                        <div class="formula-result">= {total_steel_volume_single:.8f} m³</div>
                        <div class="formula">（五）单个基础钢材重量：{total_steel_volume_single:.8f}m³ × 7850kg/m³</div>
                        <div class="formula-result">= {steel_weight_kg_single:.2f} kg</div>
                        <div class="formula-result">= {steel_weight_single:.4f} t</div>
                        <div class="formula">（六）总钢材重量：{steel_weight_kg_single:.2f}kg × {foundation_count}个</div>
                        <div class="formula-result">= {steel_weight_kg:.2f} kg</div>
                        <div class="formula-result">= {steel_weight:.4f} t</div>
                    </div>
                    
                    <div class="calculation">
                        <div class="formula">七、地基承载力验算</div>
                        <div class="formula">（一）基础体积：{length}m × {width}m × {height}m = {basic_volume_single:.4f} m³</div>
                        <div class="formula">（二）基础自重：{basic_volume_single:.4f}m³ × {concrete_density}KN/m³ = {basic_weight:.2f} KN</div>
                        <div class="formula">{'（素砼密度：22KN/m³）' if is_plain_concrete else '（普通混凝土密度：25KN/m³）'}</div>
                        <div class="formula">（三）总荷载：上部荷载 {float(self._load_input.text()) if self._load_input.text() else 0}KN + 基础自重 {basic_weight:.2f}KN = {total_load:.2f} KN</div>
                        <div class="formula">（四）基底面积：{length}m × {width}m = {length*width:.4f} m²</div>
                        <div class="formula">（五）基底压力：{total_load:.2f}KN ÷ {length*width:.4f}m² = {base_pressure:.2f} kPa</div>
                        <div class="formula">（六）地基承载力：{float(self._bearing_input.text()) if self._bearing_input.text() else 0} kPa</div>
                        <div class="formula">（七）验算结果：{'✅ 地基承载力满足要求' if is_bearing_satisfied else '❌ 地基承载力不满足要求'}</div>
                    </div>
                    
                    <div class="calculation">
                        <div class="formula">八、基础防腐面积计算</div>
                        <div class="formula">（一）单个基础防腐面积</div>
                        <div class="formula">基础侧面积 = (长度 + 宽度) × 2 × 基底埋深</div>
                        <div class="formula">基底埋深 = 基础高度 - 基础高出地面高度</div>
                        <div class="formula">= {height}m - {height_above_ground}m = {depth:.4f}m</div>
                        <div class="formula">= ({length}m + {width}m) × 2 × {depth}m</div>
                        <div class="formula-result">= {anticorrosion_area_single:.4f} m²</div>
                        <div class="formula">（二）总基础防腐面积：{anticorrosion_area_single:.4f}m² × {foundation_count}个</div>
                        <div class="formula-result">= {anticorrosion_area:.4f} m²</div>
                    </div>
            </body>
            </html>"""
            
            # 设置HTML结果
            self._result_text.setHtml(result_html)
        
        except ValueError as e:
            # 处理输入值转换错误
            error_html = f"""<html>
            <head>
                <style>
                    body {{ font-family: 'Microsoft YaHei UI', 'Consolas', 'Segoe UI', sans-serif; font-size: 14px; line-height: 1.6; }}
                    .error-card {{ background: qlineargradient(x1:0, y1:0, x2:0, y2:1, stop:0 #ffebee, stop:1 #ffcdd2); border: 1px solid #f44336; border-radius: 8px; padding: 20px; margin: 15px; }}
                    .error-title {{ color: #f44336; font-size: 16px; font-weight: bold; margin-bottom: 10px; }}
                    .error-message {{ color: #c62828; margin-bottom: 8px; }}
                    .error-detail {{ color: #8e0000; font-family: 'Consolas', monospace; }}
                </style>
            </head>
            <body>
                <div class="error-card">
                    <div class="error-title">计算错误</div>
                    <div class="error-message">请确保所有输入值为有效的数值。</div>
                    <div class="error-detail">错误详情：{str(e)}</div>
                </div>
            </body>
            </html>"""
            self._result_text.setHtml(error_html)
        except Exception as e:
            # 处理其他计算错误
            error_html = f"""<html>
            <head>
                <style>
                    body {{ font-family: 'Microsoft YaHei UI', 'Consolas', 'Segoe UI', sans-serif; font-size: 14px; line-height: 1.6; }}
                    .error-card {{ background: qlineargradient(x1:0, y1:0, x2:0, y2:1, stop:0 #ffebee, stop:1 #ffcdd2); border: 1px solid #f44336; border-radius: 8px; padding: 20px; margin: 15px; }}
                    .error-title {{ color: #f44336; font-size: 16px; font-weight: bold; margin-bottom: 10px; }}
                    .error-message {{ color: #c62828; }}
                    .error-detail {{ color: #8e0000; font-family: 'Consolas', monospace; margin-top: 8px; }}
                </style>
            </head>
            <body>
                <div class="error-card">
                    <div class="error-title">计算错误</div>
                    <div class="error-message">计算过程中发生了未知错误。</div>
                    <div class="error-detail">错误详情：{str(e)}</div>
                </div>
            </body>
            </html>"""
            self._result_text.setHtml(error_html)
    
    def reset(self):
        """重置插件UI到初始状态"""
        # 清空所有输入框
        self._length_input.clear()
        self._width_input.clear()
        self._height_input.clear()
        self._height_above_ground_input.clear()
        self._cushion_input.clear()
        self._grout_input.clear()
        self._bearing_input.clear()
        self._load_input.clear()
        self._plate_length_input.clear()
        self._plate_width_input.clear()
        self._plate_thickness_input.clear()
        self._anchor_count_input.clear()
        self._anchor_diam_input.clear()
        self._anchor_length_input.clear()
        self._replacement_thickness_input.clear()
        self._replacement_width_input.clear()
        self._foundation_count_input.clear()
        self._pile_count_input.clear()
        # 清空结果文本框
        self._result_text.clear()
        # 重置是否素砼按钮状态
        self._is_plain_concrete_no_btn.setChecked(True)
        self._is_plain_concrete_yes_btn.setChecked(False)
        # 重置是否打桩按钮状态
        self._is_pile_no_btn.setChecked(True)
        self._is_pile_yes_btn.setChecked(False)
        self._pile_count_input.setEnabled(False)  # 禁用桩根数输入框
        # 重置材质选项

        self._concrete_material_combo.setCurrentIndex(0)  # 默认C20
    
    def save(self, file_path):
        """保存当前输入的参数到文件
        
        Args:
            file_path: 保存文件的路径
        """
        import json
        
        # 收集所有输入参数
        data = {
            "基础长度": self._length_input.text(),
            "基础宽度": self._width_input.text(),
            "基础高度": self._height_input.text(),
            "基础高出地面高度": self._height_above_ground_input.text(),
            "垫层厚度": self._cushion_input.text(),
            "二次灌浆厚度": self._grout_input.text(),
            "地基承载力": self._bearing_input.text(),
            "上部荷载": self._load_input.text(),
            "预埋钢板长度": self._plate_length_input.text(),
            "预埋钢板宽度": self._plate_width_input.text(),
            "预埋钢板厚度": self._plate_thickness_input.text(),
            "地脚螺栓个数": self._anchor_count_input.text(),
            "地脚螺栓直径": self._anchor_diam_input.text(),
            "地脚螺栓长度": self._anchor_length_input.text(),

            "基础混凝土材质": self._concrete_material_combo.currentText(),
            "换填厚度": self._replacement_thickness_input.text(),
            "换填宽度": self._replacement_width_input.text(),
            "是否素砼": "是" if self._is_plain_concrete_yes_btn.isChecked() else "否",
            "是否打桩": "是" if self._is_pile_yes_btn.isChecked() else "否",
            "桩根数": self._pile_count_input.text(),
            "基础个数": self._foundation_count_input.text()
        }
        
        # 将数据保存到文件
        try:
            with open(file_path, 'w', encoding='utf-8') as f:
                json.dump(data, f, ensure_ascii=False, indent=4)
            # 更新结果文本框
            self._result_text.append(f"参数已保存到: {file_path}")
        except Exception as e:
            self._result_text.append(f"保存失败: {str(e)}")
    
    def open(self, file_path):
        """从文件中加载参数并填充到输入框
        
        Args:
            file_path: 加载文件的路径
        """
        import json
        
        try:
            # 从文件中读取数据
            with open(file_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
            
            # 填充输入框
            if "基础长度" in data:
                self._length_input.setText(data["基础长度"])
            if "基础宽度" in data:
                self._width_input.setText(data["基础宽度"])
            if "基础高度" in data:
                self._height_input.setText(data["基础高度"])
            if "基础高出地面高度" in data:
                self._height_above_ground_input.setText(data["基础高出地面高度"])
            elif "基底埋深" in data:
                # 兼容旧版数据
                self._height_above_ground_input.setText(data["基底埋深"])
            if "垫层厚度" in data:
                self._cushion_input.setText(data["垫层厚度"])
            if "二次灌浆厚度" in data:
                self._grout_input.setText(data["二次灌浆厚度"])
            if "地基承载力" in data:
                self._bearing_input.setText(data["地基承载力"])
            if "上部荷载" in data:
                self._load_input.setText(data["上部荷载"])
            if "预埋钢板长度" in data:
                self._plate_length_input.setText(data["预埋钢板长度"])
            if "预埋钢板宽度" in data:
                self._plate_width_input.setText(data["预埋钢板宽度"])
            if "预埋钢板厚度" in data:
                self._plate_thickness_input.setText(data["预埋钢板厚度"])
            if "地脚螺栓个数" in data:
                self._anchor_count_input.setText(data["地脚螺栓个数"])
            if "地脚螺栓直径" in data:
                self._anchor_diam_input.setText(data["地脚螺栓直径"])
            if "地脚螺栓长度" in data:
                self._anchor_length_input.setText(data["地脚螺栓长度"])

            if "基础混凝土材质" in data:
                material = data["基础混凝土材质"]
                index = self._concrete_material_combo.findText(material)
                if index >= 0:
                    self._concrete_material_combo.setCurrentIndex(index)
            if "换填厚度" in data:
                self._replacement_thickness_input.setText(data["换填厚度"])
            if "换填宽度" in data:
                self._replacement_width_input.setText(data["换填宽度"])
            if "是否素砼" in data:
                is_plain_concrete = data["是否素砼"]
                if is_plain_concrete == "是":
                    self._is_plain_concrete_yes_btn.setChecked(True)
                    self._is_plain_concrete_no_btn.setChecked(False)
                else:
                    self._is_plain_concrete_no_btn.setChecked(True)
                    self._is_plain_concrete_yes_btn.setChecked(False)
            # 兼容旧版数据
            elif "是否泵基础" in data:
                is_pump = data["是否泵基础"]
                # 旧版泵基础的"是"对应现在的"否"，因为泵基础通常不是素砼
                if is_pump == "是":
                    self._is_plain_concrete_no_btn.setChecked(True)
                    self._is_plain_concrete_yes_btn.setChecked(False)
                else:
                    self._is_plain_concrete_no_btn.setChecked(True)
                    self._is_plain_concrete_yes_btn.setChecked(False)
            if "是否打桩" in data:
                is_pile = data["是否打桩"]
                if is_pile == "是":
                    self._is_pile_yes_btn.setChecked(True)
                    self._is_pile_no_btn.setChecked(False)
                    self._pile_count_input.setEnabled(True)
                else:
                    self._is_pile_no_btn.setChecked(True)
                    self._is_pile_yes_btn.setChecked(False)
                    self._pile_count_input.setEnabled(False)
            if "桩根数" in data:
                self._pile_count_input.setText(data["桩根数"])
            if "基础个数" in data:
                self._foundation_count_input.setText(data["基础个数"])
            
            # 更新结果文本框
            self._result_text.append(f"从文件加载参数: {file_path}")
        except Exception as e:
            self._result_text.append(f"加载失败: {str(e)}")
    
    @Slot()
    def _on_export_triggered(self):
        """处理导出计算书按钮的点击事件
        
        将计算结果导出为Word文档
        """
        # 检查是否有计算结果
        if not self._result_text.toPlainText():
            from PySide6.QtWidgets import QMessageBox
            QMessageBox.warning(self, "警告", "请先进行计算，获取计算结果后再导出计算书！")
            return
        
        # 打开文件保存对话框
        from PySide6.QtWidgets import QFileDialog
        file_path, _ = QFileDialog.getSaveFileName(
            self,
            "保存计算书",
            "块式基础计算书.docx",
            "Word 文档 (*.docx);;所有文件 (*)"
        )
        
        if not file_path:
            return
        
        # 确保文件后缀为.docx
        if not file_path.endswith('.docx'):
            file_path += '.docx'
        
        # 使用python-docx库创建Word文档
        try:
            from docx import Document
            from docx.shared import Inches, Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.oxml.ns import qn
            
            # 创建文档
            doc = Document()
            
            # 设置文档属性
            doc.core_properties.title = "块式基础计算书"
            doc.core_properties.author = "符构工具箱"
            
            # 设置默认字体为宋体，黑色
            for style in doc.styles:
                if style.name == 'Normal':
                    style.font.name = '宋体'
                    style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                    style.font.size = Pt(12)
                    style.font.color.rgb = RGBColor(0, 0, 0)  # 黑色
                    break
            
            # 设置所有标题样式为黑色
            for style_name in ['Heading 1', 'Heading 2', 'Heading 3', 'Heading 4', 'Heading 5', 'Heading 6']:
                if style_name in doc.styles:
                    heading_style = doc.styles[style_name]
                    heading_style.font.color.rgb = RGBColor(0, 0, 0)  # 黑色
            
            # 添加标题
            title = doc.add_heading("块式基础计算书", 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_run = title.runs[0]
            title_run.font.name = '宋体'
            title_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            title_run.font.size = Pt(24)
            title_run.font.color.rgb = RGBColor(0, 0, 0)  # 黑色
            
            # 添加日期
            import datetime
            today = datetime.datetime.now().strftime("%Y年%m月%d日")
            date_paragraph = doc.add_paragraph(f"计算日期：{today}")
            date_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            date_paragraph.paragraph_format.space_after = Pt(12)
            # 设置日期字体为黑色
            for run in date_paragraph.runs:
                run.font.color.rgb = RGBColor(0, 0, 0)  # 黑色
            
            # 重新执行计算，确保获取最新结果
            self._on_calculate_clicked()
            
            # 手动收集输入参数和计算结果
            # 1. 获取输入参数
            try:
                # 基础尺寸参数
                length = float(self._length_input.text()) if self._length_input.text() else 0
                width = float(self._width_input.text()) if self._width_input.text() else 0
                height = float(self._height_input.text()) if self._height_input.text() else 0
                height_above_ground = float(self._height_above_ground_input.text()) if self._height_above_ground_input.text() else 0
                depth = height - height_above_ground
                cushion_thickness = float(self._cushion_input.text()) if self._cushion_input.text() else 0
                grout_thickness = float(self._grout_input.text()) if self._grout_input.text() else 0
                
                # 力学参数
                bearing_capacity = float(self._bearing_input.text()) if self._bearing_input.text() else 0
                upper_load = float(self._load_input.text()) if self._load_input.text() else 0
                
                # 预埋钢板参数
                plate_length = float(self._plate_length_input.text()) if self._plate_length_input.text() else 0
                plate_width = float(self._plate_width_input.text()) if self._plate_width_input.text() else 0
                plate_thickness = float(self._plate_thickness_input.text()) if self._plate_thickness_input.text() else 0
                
                # 地脚螺栓参数
                anchor_count = int(self._anchor_count_input.text()) if self._anchor_count_input.text() else 0
                anchor_diam = float(self._anchor_diam_input.text()) if self._anchor_diam_input.text() else 0
                anchor_length = float(self._anchor_length_input.text()) if self._anchor_length_input.text() else 0
                
                # 基础混凝土材质
                concrete_material = self._concrete_material_combo.currentText()
                
                # 换填参数
                replacement_thickness = float(self._replacement_thickness_input.text()) if self._replacement_thickness_input.text() else 0
                replacement_width = float(self._replacement_width_input.text()) if self._replacement_width_input.text() else 0
                
                # 其他参数
                is_plain_concrete = self._is_plain_concrete_yes_btn.isChecked()
                is_pile = self._is_pile_yes_btn.isChecked()
                pile_count = int(self._pile_count_input.text()) if self._pile_count_input.text() else 0
                foundation_count = int(self._foundation_count_input.text()) if self._foundation_count_input.text() else 1
                
                # 2. 重新计算所有结果
                # 单个基础计算
                basic_volume_single = self._logic.calculate_basic_volume(length, width, height)
                cushion_volume_single = self._logic.calculate_cushion_volume(length, width, cushion_thickness)
                replacement_volume_single = self._logic.calculate_replacement_volume(length, width, replacement_width, replacement_thickness)
                plate_volume_single = self._logic.calculate_plate_volume(plate_length, plate_width, plate_thickness)
                grout_volume_single = self._logic.calculate_grout_volume(length, width, grout_thickness)
                
                # 钢材重量计算
                anchor_bolt_volume = self._logic.calculate_anchor_bolt_volume(anchor_diam, anchor_length)
                total_steel_volume_single = anchor_bolt_volume * anchor_count + plate_volume_single
                steel_weight_kg_single = total_steel_volume_single * 7850  # 钢材密度：7850 kg/m³
                steel_weight_single = steel_weight_kg_single / 1000  # 转换为吨(t)
                
                # 地基承载力验算
                bearing_check_result_single = self._logic.check_bearing_capacity(
                    upper_load,
                    length,
                    width,
                    height,
                    bearing_capacity,
                    is_plain_concrete
                )
                is_bearing_satisfied, basic_weight, total_load, base_pressure, concrete_density = bearing_check_result_single
                
                # 基础防腐面积计算
                anticorrosion_area_single = self._logic.calculate_anticorrosion_area(length, width, depth)
                
                # 乘以基础个数，得到最终结果
                basic_volume = basic_volume_single * foundation_count
                cushion_volume = cushion_volume_single * foundation_count
                replacement_volume = replacement_volume_single * foundation_count
                steel_weight = steel_weight_single * foundation_count
                steel_weight_kg = steel_weight_kg_single * foundation_count
                grout_volume = grout_volume_single * foundation_count
                plate_volume = plate_volume_single * foundation_count
                anticorrosion_area = anticorrosion_area_single * foundation_count
                
                # 3. 创建Word文档内容
                
                # 1. 添加输入参数标题
                doc.add_heading("一、输入参数", level=1)
                
                # 添加输入参数表格
                param_doc_table = doc.add_table(rows=0, cols=3)
                param_doc_table.style = 'Table Grid'
                
                # 设置表格列宽
                for col in param_doc_table.columns:
                    col.width = Inches(2.0)
                
                # 添加输入参数行
                param_rows = [
                    [f"基础长度：{length}m", f"基础宽度：{width}m", f"基础高度：{height}m"],
                    [f"基础高出地面高度：{height_above_ground}m", f"基底埋深：{depth:.4f}m", f"二次灌浆厚度：{grout_thickness}mm"],
                    [f"地基承载力：{bearing_capacity}kPa", f"上部荷载：{upper_load}KN", f"预埋钢板：{plate_length}m × {plate_width}m × {plate_thickness}mm"],
                    [f"地脚螺栓个数：{anchor_count}个", f"地脚螺栓直径：{anchor_diam}mm", f"地脚螺栓长度：{anchor_length}mm"],
                    [f"基础混凝土材质：{concrete_material}", f"是否素砼：{'是' if is_plain_concrete else '否'}", ""],
                    [f"是否打桩：{'是' if is_pile else '否'}", f"桩根数：{pile_count}根", f"基础个数：{foundation_count}个"],
                    [f"换填厚度：{replacement_thickness}m", f"换填宽度：{replacement_width}m", ""]
                ]
                
                for row_data in param_rows:
                    row_cells = param_doc_table.add_row().cells
                    for i, cell_data in enumerate(row_data):
                        row_cells[i].text = cell_data
                
                # 2. 添加计算过程标题
                heading = doc.add_heading("二、计算过程", level=1)
                # 设置标题为黑色
                for run in heading.runs:
                    run.font.color.rgb = RGBColor(0, 0, 0)  # 黑色
                
                # 添加计算过程
                
                # 辅助函数：设置段落为黑色
                def set_paragraph_black(paragraph):
                    for run in paragraph.runs:
                        run.font.color.rgb = RGBColor(0, 0, 0)  # 黑色
                
                # 基础体积计算
                heading = doc.add_heading("1. 基础体积计算", level=2)
                set_paragraph_black(heading)
                p1 = doc.add_paragraph(f"（1）单个基础体积：{length}m × {width}m × {height}m")
                set_paragraph_black(p1)
                p2 = doc.add_paragraph(f"    = {basic_volume_single:.4f} m³")
                set_paragraph_black(p2)
                p3 = doc.add_paragraph(f"（2）总基础体积：{basic_volume_single:.4f}m³ × {foundation_count}个")
                set_paragraph_black(p3)
                p4 = doc.add_paragraph(f"    = {basic_volume:.4f} m³")
                set_paragraph_black(p4)
                
                # 垫层体积计算
                heading = doc.add_heading("2. 垫层体积计算", level=2)
                set_paragraph_black(heading)
                p1 = doc.add_paragraph(f"（1）单个基础垫层体积：({length}m + 0.1m) × ({width}m + 0.1m) × {cushion_thickness}m")
                set_paragraph_black(p1)
                p2 = doc.add_paragraph(f"    = {cushion_volume_single:.4f} m³")
                set_paragraph_black(p2)
                p3 = doc.add_paragraph(f"（2）总垫层体积：{cushion_volume_single:.4f}m³ × {foundation_count}个")
                set_paragraph_black(p3)
                p4 = doc.add_paragraph(f"    = {cushion_volume:.4f} m³")
                set_paragraph_black(p4)
                
                # 换填级配砂石体积计算
                heading = doc.add_heading("3. 换填级配砂石体积计算", level=2)
                set_paragraph_black(heading)
                p1 = doc.add_paragraph(f"（1）单个基础换填级配砂石体积：({length}m + {replacement_width}m) × ({width}m + {replacement_width}m) × {replacement_thickness}m")
                set_paragraph_black(p1)
                p2 = doc.add_paragraph(f"    = {replacement_volume_single:.4f} m³")
                set_paragraph_black(p2)
                p3 = doc.add_paragraph(f"（2）总换填级配砂石体积：{replacement_volume_single:.4f}m³ × {foundation_count}个")
                set_paragraph_black(p3)
                p4 = doc.add_paragraph(f"    = {replacement_volume:.4f} m³")
                set_paragraph_black(p4)
                
                # 二次灌浆体积计算
                heading = doc.add_heading("4. 二次灌浆体积计算", level=2)
                set_paragraph_black(heading)
                p1 = doc.add_paragraph(f"（1）单个基础二次灌浆体积：{length}m × {width}m × {grout_thickness}mm")
                set_paragraph_black(p1)
                p2 = doc.add_paragraph(f"    = {length}m × {width}m × {grout_thickness/1000:.6f}m")
                set_paragraph_black(p2)
                p3 = doc.add_paragraph(f"    = {grout_volume_single:.4f} m³")
                set_paragraph_black(p3)
                p4 = doc.add_paragraph(f"（2）总二次灌浆体积：{grout_volume_single:.4f}m³ × {foundation_count}个")
                set_paragraph_black(p4)
                p5 = doc.add_paragraph(f"    = {grout_volume:.4f} m³")
                set_paragraph_black(p5)
                
                # 预埋钢板体积计算
                heading = doc.add_heading("5. 预埋钢板体积计算", level=2)
                set_paragraph_black(heading)
                p1 = doc.add_paragraph(f"（1）单个预埋钢板体积：{plate_length}m × {plate_width}m × {plate_thickness}mm")
                set_paragraph_black(p1)
                p2 = doc.add_paragraph(f"    = {plate_length}m × {plate_width}m × {plate_thickness/1000:.6f}m")
                set_paragraph_black(p2)
                p3 = doc.add_paragraph(f"    = {plate_volume_single:.6f} m³")
                set_paragraph_black(p3)
                p4 = doc.add_paragraph(f"（2）总预埋钢板体积：{plate_volume_single:.6f}m³ × {foundation_count}个")
                set_paragraph_black(p4)
                p5 = doc.add_paragraph(f"    = {plate_volume:.4f} m³")
                set_paragraph_black(p5)
                
                # 钢材重量计算
                heading = doc.add_heading("6. 钢材重量计算", level=2)
                set_paragraph_black(heading)
                p1 = doc.add_paragraph(f"（1）单根地脚螺栓体积：π × ({anchor_diam/1000/2:.4f}m)² × {anchor_length/1000:.4f}m")
                set_paragraph_black(p1)
                p2 = doc.add_paragraph(f"    = {anchor_bolt_volume:.8f} m³")
                set_paragraph_black(p2)
                p3 = doc.add_paragraph(f"（2）单根地脚螺栓总重量：{anchor_bolt_volume:.8f}m³ × {anchor_count}根")
                set_paragraph_black(p3)
                p4 = doc.add_paragraph(f"    = {anchor_bolt_volume * anchor_count:.8f} m³")
                set_paragraph_black(p4)
                p5 = doc.add_paragraph(f"（3）预埋钢板体积：{plate_length}m × {plate_width}m × {plate_thickness/1000:.6f}m")
                set_paragraph_black(p5)
                p6 = doc.add_paragraph(f"    = {plate_volume_single:.6f} m³")
                set_paragraph_black(p6)
                p7 = doc.add_paragraph(f"（4）单个基础总钢材体积：地脚螺栓体积 + 预埋钢板体积")
                set_paragraph_black(p7)
                p8 = doc.add_paragraph(f"    = {anchor_bolt_volume * anchor_count:.8f}m³ + {plate_volume_single:.6f}m³")
                set_paragraph_black(p8)
                p9 = doc.add_paragraph(f"    = {total_steel_volume_single:.8f} m³")
                set_paragraph_black(p9)
                p10 = doc.add_paragraph(f"（5）单个基础钢材重量：{total_steel_volume_single:.8f}m³ × 7850kg/m³")
                set_paragraph_black(p10)
                p11 = doc.add_paragraph(f"    = {steel_weight_kg_single:.2f} kg")
                set_paragraph_black(p11)
                p12 = doc.add_paragraph(f"    = {steel_weight_single:.4f} t")
                set_paragraph_black(p12)
                p13 = doc.add_paragraph(f"（6）总钢材重量：{steel_weight_kg_single:.2f}kg × {foundation_count}个")
                set_paragraph_black(p13)
                p14 = doc.add_paragraph(f"    = {steel_weight_kg:.2f} kg")
                set_paragraph_black(p14)
                p15 = doc.add_paragraph(f"    = {steel_weight:.4f} t")
                set_paragraph_black(p15)
                
                # 地基承载力验算
                heading = doc.add_heading("7. 地基承载力验算", level=2)
                set_paragraph_black(heading)
                p1 = doc.add_paragraph(f"（1）基础体积：{length}m × {width}m × {height}m = {basic_volume_single:.4f} m³")
                set_paragraph_black(p1)
                p2 = doc.add_paragraph(f"（2）基础自重：{basic_volume_single:.4f}m³ × {concrete_density}KN/m³ = {basic_weight:.2f} KN")
                set_paragraph_black(p2)
                p3 = doc.add_paragraph(f"    {'（素砼密度：22KN/m³）' if is_plain_concrete else '（普通混凝土密度：25KN/m³）'}")
                set_paragraph_black(p3)
                p4 = doc.add_paragraph(f"（3）总荷载：上部荷载 {upper_load}KN + 基础自重 {basic_weight:.2f}KN = {total_load:.2f} KN")
                set_paragraph_black(p4)
                p5 = doc.add_paragraph(f"（4）基底面积：{length}m × {width}m = {length*width:.4f} m²")
                set_paragraph_black(p5)
                p6 = doc.add_paragraph(f"（5）基底压力：{total_load:.2f}KN ÷ {length*width:.4f}m² = {base_pressure:.2f} kPa")
                set_paragraph_black(p6)
                p7 = doc.add_paragraph(f"（6）地基承载力：{bearing_capacity} kPa")
                set_paragraph_black(p7)
                p8 = doc.add_paragraph(f"（7）验算结果：{'地基承载力满足要求' if is_bearing_satisfied else '地基承载力不满足要求'}")
                set_paragraph_black(p8)
                
                # 基础防腐面积计算
                heading = doc.add_heading("8. 基础防腐面积计算", level=2)
                set_paragraph_black(heading)
                p1 = doc.add_paragraph(f"（1）单个基础防腐面积")
                set_paragraph_black(p1)
                p2 = doc.add_paragraph(f"    基础侧面积 = (长度 + 宽度) × 2 × 基底埋深")
                set_paragraph_black(p2)
                p3 = doc.add_paragraph(f"    基底埋深 = 基础高度 - 基础高出地面高度")
                set_paragraph_black(p3)
                p4 = doc.add_paragraph(f"    = {height}m - {height_above_ground}m = {depth:.4f}m")
                set_paragraph_black(p4)
                p5 = doc.add_paragraph(f"    = ({length}m + {width}m) × 2 × {depth}m")
                set_paragraph_black(p5)
                p6 = doc.add_paragraph(f"    = {anticorrosion_area_single:.4f} m²")
                set_paragraph_black(p6)
                p7 = doc.add_paragraph(f"（2）总基础防腐面积：{anticorrosion_area_single:.4f}m² × {foundation_count}个")
                set_paragraph_black(p7)
                p8 = doc.add_paragraph(f"    = {anticorrosion_area:.4f} m²")
                set_paragraph_black(p8)
                
                # 3. 添加计算结果标题
                heading = doc.add_heading("三、最终计算结果", level=1)
                set_paragraph_black(heading)
                
                # 添加表格显示最终结果
                final_results_table = doc.add_table(rows=1, cols=2)
                final_results_table.style = 'Table Grid'
                
                # 设置表格列宽
                final_results_table.columns[0].width = Inches(3.0)
                final_results_table.columns[1].width = Inches(2.0)
                
                # 添加表格标题行
                hdr_cells = final_results_table.rows[0].cells
                hdr_cells[0].text = '项目'
                hdr_cells[1].text = '结果'
                
                # 设置标题行样式
                for cell in hdr_cells:
                    cell.paragraphs[0].runs[0].font.bold = True
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    # 设置标题行文字为黑色
                    for run in cell.paragraphs[0].runs:
                        run.font.color.rgb = RGBColor(0, 0, 0)  # 黑色
                
                # 添加最终结果数据
                final_results = [
                    ["基础体积", f"{basic_volume:.4f} m³"],
                    ["垫层体积", f"{cushion_volume:.4f} m³"],
                    ["换填级配砂石体积", f"{replacement_volume:.4f} m³"],
                    ["二次灌浆体积", f"{grout_volume:.4f} m³"],
                    ["钢材重量", f"{steel_weight:.4f} t"],
                    ["基础防腐面积", f"{anticorrosion_area:.4f} m²"],
                    ["地基承载力验算", f"{'满足要求' if is_bearing_satisfied else '不满足要求'}"]
                ]
                
                for result_data in final_results:
                    row_cells = final_results_table.add_row().cells
                    row_cells[0].text = result_data[0]
                    row_cells[1].text = result_data[1]
                    row_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    # 设置表格单元格文字为黑色
                    for cell in row_cells:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.color.rgb = RGBColor(0, 0, 0)  # 黑色
                
                # 添加页脚
                footer = doc.sections[0].footer
                footer_paragraph = footer.paragraphs[0]
                footer_paragraph.text = "符构工具箱 - 块式基础计算"
                footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                # 设置页脚文字为黑色
                for run in footer_paragraph.runs:
                    run.font.color.rgb = RGBColor(0, 0, 0)  # 黑色
                
                # 保存文档
                doc.save(file_path)
                
                from PySide6.QtWidgets import QMessageBox
                from PySide6.QtGui import QDesktopServices
                from PySide6.QtCore import QUrl
                
                msg_box = QMessageBox()
                msg_box.setWindowTitle("成功")
                msg_box.setText(f"计算书已成功保存到：\n{file_path}")
                msg_box.setIcon(QMessageBox.Information)
                msg_box.addButton("确定", QMessageBox.AcceptRole)
                msg_box.addButton("打开文件", QMessageBox.ActionRole)
                
                result = msg_box.exec_()
                if result == 1:  # 打开文件按钮被点击
                    QDesktopServices.openUrl(QUrl.fromLocalFile(file_path))
                
            except ValueError as e:
                from PySide6.QtWidgets import QMessageBox
                QMessageBox.critical(self, "错误", f"导出计算书失败：\n输入参数错误：{str(e)}")
            except Exception as e:
                from PySide6.QtWidgets import QMessageBox
                QMessageBox.critical(self, "错误", f"导出计算书失败：\n计算过程错误：{str(e)}")
                import traceback
                traceback.print_exc()
        except Exception as e:
            from PySide6.QtWidgets import QMessageBox
            QMessageBox.critical(self, "错误", f"导出计算书失败：\n{str(e)}")
            import traceback
            traceback.print_exc()

    @Slot()
    def _on_export_material_triggered(self):
        """处理导出料表按钮的点击事件"""
        from PySide6.QtWidgets import QFileDialog, QMessageBox
        from datetime import datetime
        try:
            # 检查是否有计算结果
            if not self._result_text.toPlainText():
                QMessageBox.warning(self, "警告", "请先进行计算，获取计算结果后再导出料表！")
                return

            length = float(self._length_input.text()) if self._length_input.text() else 0
            width = float(self._width_input.text()) if self._width_input.text() else 0
            height = float(self._height_input.text()) if self._height_input.text() else 0
            height_above_ground = float(self._height_above_ground_input.text()) if self._height_above_ground_input.text() else 0
            depth = height - height_above_ground
            foundation_count = int(self._foundation_count_input.text()) if self._foundation_count_input.text() else 1
            cushion_thickness = float(self._cushion_input.text()) if self._cushion_input.text() else 0
            replacement_width = float(self._replacement_width_input.text()) if self._replacement_width_input.text() else 0
            replacement_thickness = float(self._replacement_thickness_input.text()) if self._replacement_thickness_input.text() else 0
            grout_thickness = float(self._grout_input.text()) if self._grout_input.text() else 0
            is_plain_concrete = self._is_plain_concrete_yes_btn.isChecked()
            is_pile = self._is_pile_yes_btn.isChecked()

            basic_volume_single = self._logic.calculate_basic_volume(length, width, height)
            cushion_volume_single = self._logic.calculate_cushion_volume(length, width, cushion_thickness)
            replacement_volume_single = self._logic.calculate_replacement_volume(length, width, replacement_width, replacement_thickness)
            grout_volume_single = self._logic.calculate_grout_volume(length, width, grout_thickness)
            anticorrosion_area_single = self._logic.calculate_anticorrosion_area(length, width, depth)

            basic_volume = basic_volume_single * foundation_count
            cushion_volume = cushion_volume_single * foundation_count
            replacement_volume = replacement_volume_single * foundation_count
            grout_volume = grout_volume_single * foundation_count
            anticorrosion_area = anticorrosion_area_single * foundation_count

            # 获取预埋钢板参数
            plate_length = float(self._plate_length_input.text()) if self._plate_length_input.text() else 0
            plate_width = float(self._plate_width_input.text()) if self._plate_width_input.text() else 0
            plate_thickness = float(self._plate_thickness_input.text()) if self._plate_thickness_input.text() else 0
            
            # 计算预埋钢板体积（单个基础）
            plate_volume_single = self._logic.calculate_plate_volume(plate_length, plate_width, plate_thickness)
            plate_volume = plate_volume_single * foundation_count
            
            # 获取地脚螺栓相关参数
            anchor_diam = float(self._anchor_diam_input.text()) if self._anchor_diam_input.text() else 0
            anchor_count = int(self._anchor_count_input.text()) if self._anchor_count_input.text() else 0
            anchor_length = float(self._anchor_length_input.text()) if self._anchor_length_input.text() else 0
            
            # 计算钢材重量（单个基础）
            anchor_bolt_volume = self._logic.calculate_anchor_bolt_volume(anchor_diam, anchor_length)
            # 总钢材体积 = 地脚螺栓体积 + 预埋钢板体积
            total_steel_volume_single = anchor_bolt_volume * anchor_count + plate_volume_single
            steel_weight_kg_single = total_steel_volume_single * 7850  # 钢材密度：7850 kg/m³
            steel_weight_single = steel_weight_kg_single / 1000  # 转换为吨(t)
            
            steel_weight = steel_weight_single * foundation_count
            steel_weight_kg = steel_weight_kg_single * foundation_count
            
            # 获取桩根数
            pile_count = int(self._pile_count_input.text()) if self._pile_count_input.text() else 0

            concrete_material = self._concrete_material_combo.currentText()
            concrete_type = "素砼" if is_plain_concrete else "钢筋砼"
            first_item_name = f"{concrete_material}{concrete_type}基础体积"

            file_path, _ = QFileDialog.getSaveFileName(
                self,
                "导出料表",
                f"块式基础料表_{datetime.now().strftime('%Y%m%d')}.xlsx",
                "Excel Files (*.xlsx)"
            )

            if file_path:
                from openpyxl import Workbook
                from openpyxl.styles import Font, Alignment, Border, Side

                wb = Workbook()
                ws = wb.active
                ws.title = "料表"

                thin_border = Border(
                    left=Side(style='thin'),
                    right=Side(style='thin'),
                    top=Side(style='thin'),
                    bottom=Side(style='thin')
                )

                header_font = Font(bold=True, size=12)
                center_alignment = Alignment(horizontal='center', vertical='center')

                data = [
                    [f"{first_item_name} (m³)", round(basic_volume, 3)],
                    ["垫层体积 (m³)", round(cushion_volume, 3)],
                    ["换填级配砂石体积 (m³)", round(replacement_volume, 3)],
                    ["二次灌浆料体积 (m³)", round(grout_volume, 3)],
                    ["钢材重量 (t)", round(steel_weight, 3)],
                    ["基础防腐面积 (m²)", round(anticorrosion_area, 3)]
                ]
                
                # 如果桩根数有输入值，添加桩根数这一项
                if pile_count > 0:
                    data.append(["桩根数 (根)", pile_count])

                ws['A1'] = "项目"
                ws['B1'] = "数值"
                ws['A1'].font = header_font
                ws['B1'].font = header_font
                ws['A1'].alignment = center_alignment
                ws['B1'].alignment = center_alignment
                ws['A1'].border = thin_border
                ws['B1'].border = thin_border
                ws.column_dimensions['A'].width = 30
                ws.column_dimensions['B'].width = 15

                from openpyxl.styles import NumberFormatDescriptor

                for idx, (item, value) in enumerate(data, start=2):
                    ws[f'A{idx}'] = item
                    ws[f'B{idx}'] = value
                    ws[f'A{idx}'].alignment = center_alignment
                    ws[f'B{idx}'].alignment = center_alignment
                    ws[f'A{idx}'].border = thin_border
                    ws[f'B{idx}'].border = thin_border
                    # 设置第二列为数值格式，保留三位小数
                    ws[f'B{idx}'].number_format = '0.000'

                wb.save(file_path)

                from PySide6.QtGui import QDesktopServices
                from PySide6.QtCore import QUrl
                
                msg_box = QMessageBox()
                msg_box.setWindowTitle("成功")
                msg_box.setText(f"料表已成功保存到：\n{file_path}")
                msg_box.setIcon(QMessageBox.Information)
                msg_box.addButton("确定", QMessageBox.AcceptRole)
                msg_box.addButton("打开文件", QMessageBox.ActionRole)
                
                result = msg_box.exec_()
                if result == 1:  # 打开文件按钮被点击
                    QDesktopServices.openUrl(QUrl.fromLocalFile(file_path))

        except ValueError as e:
            QMessageBox.critical(self, "错误", f"导出料表失败：\n输入参数错误：{str(e)}")
        except Exception as e:
            QMessageBox.critical(self, "错误", f"导出料表失败：\n计算过程错误：{str(e)}")
            import traceback
            traceback.print_exc()
