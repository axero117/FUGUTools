"""水池计算插件UI。"""

from __future__ import annotations

import datetime
import json
import math
import os
from dataclasses import asdict

from datetime import datetime
from docx import Document
from docx.shared import RGBColor
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from PySide6.QtWidgets import QMessageBox, QFileDialog
from PySide6.QtGui import QDesktopServices
from PySide6.QtCore import QUrl

from PySide6.QtCore import Qt, Slot
from PySide6.QtGui import QDoubleValidator, QIntValidator
from PySide6.QtWidgets import (
    QComboBox,
    QGridLayout,
    QGroupBox,
    QHBoxLayout,
    QLabel,
    QLineEdit,
    QMessageBox,
    QPushButton,
    QScrollArea,
    QTextEdit,
    QVBoxLayout,
    QWidget,
)

from plugins.Pool_Tank_Foundation.logic import (
    PoolFoundationInput,
    PoolTankFoundationLogic,
)


class PoolTankFoundationWidget(QWidget):
    """水池计算UI组件。"""

    def __init__(self, parent=None):
        super().__init__(parent)
        self.setObjectName("poolTankFoundationWidget")

        self._logic = PoolTankFoundationLogic()
        self._inputs: dict[str, QLineEdit] = {}
        self._last_result = None
        self._last_payload = None

        self._init_ui()
        self._apply_compact_input_text_style()
        self._setup_validators()
        self._connect_signals()
        self._on_balancing_layer_changed()
        self._on_column_height_type_changed()

    def _apply_compact_input_text_style(self):
        """统一输入控件文字大小，避免文字显示大于输入框高度。"""
        # 仅针对本模块生效，避免影响其他插件
        self.setStyleSheet(
            """
            #poolTankFoundationWidget QLineEdit {
                font-size: 12px;
                padding: 1px 4px;
            }
            #poolTankFoundationWidget QComboBox {
                font-size: 12px;
                min-height: 24px;
                padding: 1px 4px;
            }
            #poolTankFoundationWidget QLabel {
                font-size: 12px;
            }
            """
        )

    def _init_ui(self):
        """初始化UI布局。"""
        root_layout = QVBoxLayout(self)
        root_layout.setContentsMargins(6, 6, 6, 6)
        root_layout.setSpacing(6)

        scroll = QScrollArea()
        scroll.setWidgetResizable(True)
        root_layout.addWidget(scroll)

        container = QWidget()
        scroll.setWidget(container)
        main_layout = QVBoxLayout(container)
        main_layout.setContentsMargins(6, 6, 6, 6)
        main_layout.setSpacing(6)

        title = QLabel("水池计算")
        title.setObjectName("moduleTitle")
        title.setAlignment(Qt.AlignCenter)
        main_layout.addWidget(title)

        # 参数区：2x2紧凑布局，尽量在同一屏显示完整
        params_grid = QGridLayout()
        params_grid.setContentsMargins(0, 0, 0, 0)
        params_grid.setHorizontalSpacing(8)
        params_grid.setVerticalSpacing(6)

        g1 = self._build_geometry_group()
        g2 = self._build_balancing_group()
        g3 = self._build_internal_member_group()
        g4 = self._build_anti_floating_group()

        for g in (g1, g2, g3):
            g.setMaximumHeight(260)
        g4.setMaximumHeight(320)

        params_grid.addWidget(g1, 0, 0)
        params_grid.addWidget(g2, 0, 1)
        params_grid.addWidget(g3, 1, 0)
        params_grid.addWidget(g4, 1, 1)
        main_layout.addLayout(params_grid)

        # 操作按钮
        btn_layout = QHBoxLayout()
        self._btn_calc = QPushButton("计算")
        self._btn_calc.setProperty("moduleAction", "primary")
        self._btn_reset = QPushButton("重置")
        self._btn_reset.setProperty("moduleAction", "ghost")
        self._btn_export = QPushButton("导出计算书")
        self._btn_export.setProperty("moduleAction", "secondary")
        btn_layout.addWidget(self._btn_calc)
        btn_layout.addWidget(self._btn_reset)
        btn_layout.addWidget(self._btn_export)
        btn_layout.addStretch()
        main_layout.addLayout(btn_layout)

        # 输出区
        result_group = QGroupBox("计算结果")
        result_group.setProperty("moduleCard", True)
        result_layout = QVBoxLayout(result_group)
        self._result = QTextEdit()
        self._result.setReadOnly(True)
        self._result.setMinimumHeight(320)
        result_layout.addWidget(self._result)
        main_layout.addWidget(result_group)

    def _build_geometry_group(self) -> QGroupBox:
        group = QGroupBox("一、几何参数")
        group.setProperty("moduleCard", True)
        grid = QGridLayout(group)

        self._add_line(grid, 0, "内长", "L", "m")
        self._add_line(grid, 1, "内宽", "W", "m")
        self._add_line(grid, 2, "池壁总高度", "H_total", "m")
        self._add_line(grid, 3, "高出地面高度", "H_out", "m")
        self._add_line(grid, 4, "顶板厚度", "h_slab", "m")
        self._add_line(grid, 5, "池壁厚度", "t_wall", "m")
        self._add_line(grid, 6, "底板厚度", "h_footing", "m")
        self._add_line(grid, 7, "底板外挑宽度", "b_overhang", "m")
        return group

    def _build_balancing_group(self) -> QGroupBox:
        group = QGroupBox("二、平衡层与换填参数")
        group.setProperty("moduleCard", True)
        grid = QGridLayout(group)

        grid.addWidget(QLabel("有无平衡层"), 0, 0)
        self._has_balancing_layer = QComboBox()
        self._has_balancing_layer.addItems(["有", "无"])
        grid.addWidget(self._has_balancing_layer, 0, 1)
        grid.addWidget(QLabel(""), 0, 2)

        self._add_line(grid, 1, "平衡层宽出底板距离", "c_bal_overhang", "m")
        self._add_line(grid, 2, "平衡层厚度", "h_bal", "m")
        self._add_line(grid, 3, "平衡层材料重度", "gamma_bal", "kN/m³", "18")
        self._add_line(grid, 4, "换填砂石厚度", "h_cushion", "m")
        self._add_line(grid, 5, "换填宽出距离", "d_cushion_overhang", "m")
        self._add_line(grid, 6, "换填砂石重度", "gamma_cushion", "kN/m³", "19")
        self._add_line(grid, 7, "垫层厚度", "h_mat", "m", "0")
        return group

    def _build_internal_member_group(self) -> QGroupBox:
        group = QGroupBox("三、内部构件参数")
        group.setProperty("moduleCard", True)
        grid = QGridLayout(group)

        self._add_line(grid, 0, "隔墙总混凝土体积", "V_partition_total", "m³", "0")
        self._add_line(grid, 1, "柱截面长", "column_a", "m", "0")
        self._add_line(grid, 2, "柱截面宽", "column_b", "m", "0")

        grid.addWidget(QLabel("柱高类型"), 3, 0)
        self._column_height_type = QComboBox()
        self._column_height_type.addItems(["内净高", "自定义"])
        grid.addWidget(self._column_height_type, 3, 1)
        grid.addWidget(QLabel(""), 3, 2)

        self._add_line(grid, 4, "自定义柱高", "column_height_custom", "m", "0")

        grid.addWidget(QLabel("柱个数"), 5, 0)
        self._column_count = QLineEdit("0")
        self._column_count.setFixedHeight(24)
        grid.addWidget(self._column_count, 5, 1)
        grid.addWidget(QLabel("个"), 5, 2)

        return group

    def _build_anti_floating_group(self) -> QGroupBox:
        group = QGroupBox("四、抗浮参数")
        group.setProperty("moduleCard", True)
        grid = QGridLayout(group)

        self._add_line(grid, 0, "地下水位深度（地面以下）", "water_table_elev", "m", "0")
        self._inputs["water_table_elev"].setPlaceholderText("例如：0（平地面），1.2（地面下1.2m）")
        self._add_line(grid, 1, "水重度", "gamma_water", "kN/m³", "10")
        self._add_line(grid, 2, "混凝土重度", "gamma_concrete", "kN/m³", "25")
        self._add_line(grid, 3, "池内实际水深", "water_depth_inner", "m", "0")
        self._add_line(grid, 4, "填土天然重度", "gamma_soil_natural", "kN/m³", "18")
        self._add_line(grid, 5, "填土饱和重度", "gamma_soil_sat", "kN/m³", "18")
        self._add_line(grid, 6, "顶板活荷载", "roof_live_load", "kN/m²", "0")
        self._add_line(grid, 7, "地面活荷载", "ground_live_load", "kN/m²", "0")
        return group

    def _add_line(
        self,
        grid: QGridLayout,
        row: int,
        label: str,
        key: str,
        unit: str,
        default: str = "",
    ):
        """在网格中添加一行“标签-输入-单位”，并缓存输入框引用。"""
        grid.addWidget(QLabel(label), row, 0)
        edit = QLineEdit(default)
        edit.setFixedHeight(24)
        self._inputs[key] = edit
        grid.addWidget(edit, row, 1)
        grid.addWidget(QLabel(unit), row, 2)

    def _setup_validators(self):
        """设置输入校验器。"""
        float_validator = QDoubleValidator(0.0, 1e9, 6, self)
        int_validator = QIntValidator(0, 1000000, self)

        for edit in self._inputs.values():
            edit.setValidator(float_validator)

        self._column_count.setValidator(int_validator)

    def _connect_signals(self):
        self._btn_calc.clicked.connect(self._on_calculate)
        self._btn_reset.clicked.connect(self.reset)
        self._btn_export.clicked.connect(self._on_export_triggered)
        self._has_balancing_layer.currentIndexChanged.connect(self._on_balancing_layer_changed)
        self._column_height_type.currentIndexChanged.connect(self._on_column_height_type_changed)

    @Slot()
    def _on_balancing_layer_changed(self):
        """根据是否有平衡层，切换相关输入启用状态。"""
        has_bal = self._has_balancing_layer.currentText() == "有"
        self._inputs["c_bal_overhang"].setEnabled(has_bal)
        self._inputs["h_bal"].setEnabled(has_bal)
        self._inputs["gamma_bal"].setEnabled(has_bal)

        self._apply_disabled_emphasis(self._inputs["c_bal_overhang"], not has_bal)
        self._apply_disabled_emphasis(self._inputs["h_bal"], not has_bal)
        self._apply_disabled_emphasis(self._inputs["gamma_bal"], not has_bal)

    @Slot()
    def _on_column_height_type_changed(self):
        """根据柱高类型，控制自定义柱高输入框。"""
        is_custom = self._column_height_type.currentText() == "自定义"
        self._inputs["column_height_custom"].setEnabled(is_custom)
        self._apply_disabled_emphasis(self._inputs["column_height_custom"], not is_custom)

    @staticmethod
    def _apply_disabled_emphasis(edit: QLineEdit, disabled: bool):
        """禁用输入框时加深底色，让状态更直观。"""
        if disabled:
            edit.setStyleSheet(
                "QLineEdit { background-color: #d7dbe2; color: #616975; border: 1px solid #b2bac6; }"
            )
        else:
            edit.setStyleSheet("")

    def _to_float(self, key: str, default: float = 0.0) -> float:
        text = self._inputs[key].text().strip()
        return float(text) if text else default

    def _collect_input(self) -> PoolFoundationInput:
        """收集并构造计算输入对象。"""
        # 约定：
        # 1) 地下水位输入为“地面以下深度”（m，向下为正）；
        # 2) 底板底标高由几何自动计算，不再手动输入：
        #    底板底埋深 = 池壁总高度 - 水池高出地面高度 + 底板厚度
        #    底板底标高 = -底板底埋深
        water_table_depth = self._to_float("water_table_elev", 0.0)
        H_total = self._to_float("H_total")
        H_out = self._to_float("H_out")
        h_footing = self._to_float("h_footing")

        bottom_depth = H_total - H_out + h_footing
        bottom_elev_auto = -bottom_depth
        water_table_elev_converted = -water_table_depth

        return PoolFoundationInput(
            # 几何
            L=self._to_float("L"),
            W=self._to_float("W"),
            H_total=H_total,
            H_out=H_out,
            h_slab=self._to_float("h_slab"),
            t_wall=self._to_float("t_wall"),
            h_footing=h_footing,
            b_overhang=self._to_float("b_overhang"),
            # 平衡层与换填
            has_balancing_layer=self._has_balancing_layer.currentText() == "有",
            c_bal_overhang=self._to_float("c_bal_overhang"),
            h_bal=self._to_float("h_bal"),
            gamma_bal=self._to_float("gamma_bal", 18.0),
            h_cushion=self._to_float("h_cushion"),
            d_cushion_overhang=self._to_float("d_cushion_overhang"),
            gamma_cushion=self._to_float("gamma_cushion", 19.0),
            h_mat=self._to_float("h_mat", 0.0),
            # 内部构件
            V_partition_total=self._to_float("V_partition_total", 0.0),
            column_a=self._to_float("column_a", 0.0),
            column_b=self._to_float("column_b", 0.0),
            column_height_type=self._column_height_type.currentText(),
            column_height_custom=self._to_float("column_height_custom", 0.0),
            column_count=int(self._column_count.text().strip() or 0),
            # 抗浮
            water_table_elev=water_table_elev_converted,
            bottom_elev=bottom_elev_auto,
            gamma_water=self._to_float("gamma_water", 10.0),
            gamma_concrete=self._to_float("gamma_concrete", 25.0),
            water_depth_inner=self._to_float("water_depth_inner", 0.0),
            gamma_soil_natural=self._to_float("gamma_soil_natural", 18.0),
            gamma_soil_sat=self._to_float("gamma_soil_sat", 18.0),
            roof_live_load=self._to_float("roof_live_load", 0.0),
            ground_live_load=self._to_float("ground_live_load", 0.0),
        )

    @Slot()
    def _on_calculate(self):
        """执行计算并输出结果。"""
        try:
            payload = self._collect_input()
            result = self._logic.calculate(payload)
            self._result.setHtml(self._format_result(result))
            # 存储计算结果供导出使用
            self._last_result = result
            self._last_payload = payload
        except ValueError as e:
            QMessageBox.warning(self, "输入错误", str(e))
        except Exception as e:
            QMessageBox.critical(self, "计算失败", f"发生未知错误：{e}")

    def _format_result(self, r: dict) -> str:
        """格式化结果HTML：仅保留“最终计算结果、输入参数、计算过程”三部分。"""
        g = r["geometry"]
        v = r["volumes"]
        w = r["weights"]
        c = r["checks"]

        # 从界面取值用于“输入参数”展示（保持与输入框中文语义一致）
        L = self._to_float("L")
        W = self._to_float("W")
        H_total = self._to_float("H_total")
        H_out = self._to_float("H_out")
        h_slab = self._to_float("h_slab")
        t_wall = self._to_float("t_wall")
        h_footing = self._to_float("h_footing")
        b_overhang = self._to_float("b_overhang")
        c_bal_overhang = self._to_float("c_bal_overhang")
        h_bal = self._to_float("h_bal")
        gamma_bal = self._to_float("gamma_bal", 18.0)
        h_cushion = self._to_float("h_cushion")
        d_cushion_overhang = self._to_float("d_cushion_overhang")
        gamma_cushion = self._to_float("gamma_cushion", 19.0)
        V_partition_total = self._to_float("V_partition_total", 0.0)
        column_a = self._to_float("column_a", 0.0)
        column_b = self._to_float("column_b", 0.0)
        column_height_custom = self._to_float("column_height_custom", 0.0)
        water_depth_inner = self._to_float("water_depth_inner", 0.0)
        gamma_water = self._to_float("gamma_water", 10.0)
        gamma_concrete = self._to_float("gamma_concrete", 25.0)
        gamma_soil_natural = self._to_float("gamma_soil_natural", 18.0)
        gamma_soil_sat = self._to_float("gamma_soil_sat", 18.0)
        roof_live_load = self._to_float("roof_live_load", 0.0)
        ground_live_load = self._to_float("ground_live_load", 0.0)
        water_table_depth = self._to_float("water_table_elev", 0.0)
        column_count = int(self._column_count.text().strip() or 0)
        has_bal = self._has_balancing_layer.currentText() == "有"
        column_height_type = self._column_height_type.currentText()

        bottom_depth = H_total - H_out + h_footing
        bottom_elev_auto = -bottom_depth
        water_table_elev_converted = -water_table_depth

        def f3(x: float) -> str:
            return f"{x:.3f}"

        def ok(flag: bool) -> str:
            return "✅ 满足" if flag else "❌ 不满足"

        f_buoy = w["F_buoy"]
        if math.isclose(f_buoy, 0.0, abs_tol=1e-12):
            k_full_formula = "∞"
            k_actual_formula = "∞"
            k_empty_formula = "∞"
        else:
            k_full_formula = f3((w["G_total"] + w["G_water_full"]) / f_buoy)
            k_actual_formula = f3((w["G_total"] + w["G_water_actual"]) / f_buoy)
            k_empty_formula = f3(w["G_total"] / f_buoy)

        suggestion_text = "；".join(r["suggestions"]) if r["suggestions"] else "无"

        h_column_display = g["H_column"]
        h_column_formula = (
            f"柱高 = 内部净高 = {f3(g['H_in'])}"
            if column_height_type == "内净高"
            else f"柱高 = 自定义柱高 = {f3(column_height_custom)}"
        )

        bal_formula = (
            f"平衡层长度 = 底板长度 + 2 × 平衡层宽出底板距离 = {f3(g['L_base'])} + 2×{f3(c_bal_overhang)} = {f3(g['L_bal'])}"
            if has_bal
            else "无平衡层：平衡层长度 = 底板长度"
        )
        bal_formula_2 = (
            f"平衡层宽度 = 底板宽度 + 2 × 平衡层宽出底板距离 = {f3(g['W_base'])} + 2×{f3(c_bal_overhang)} = {f3(g['W_bal'])}"
            if has_bal
            else "无平衡层：平衡层宽度 = 底板宽度"
        )
        bal_volume_formula = (
            f"平衡层体积 = 平衡层长度 × 平衡层宽度 × 平衡层厚度 = {f3(g['L_bal'])} × {f3(g['W_bal'])} × {f3(h_bal)} = {f3(v['V_bal'])} m³"
            if has_bal
            else "无平衡层：平衡层体积 = 0"
        )

        return f"""
        <html>
        <head>
            <style>
                body {{ font-family: 'Microsoft YaHei UI', 'Consolas', 'Segoe UI', sans-serif; font-size: 13px; line-height: 1.65; background: transparent; }}
                div {{ background: transparent; }}
                .card {{ background: transparent; border: 1px solid #ccc; border-radius: 8px; padding: 14px 16px; margin-bottom: 12px; }}
                h2 {{ color: #2471a3; margin-top: 0; margin-bottom: 10px; font-size: 14px; font-weight: bold; border-bottom: 2px solid #2471a3; padding-bottom: 4px; }}
                .final-results {{ background: transparent; border: 1px solid #dee2e6; border-radius: 8px; padding: 12px; margin-bottom: 12px; }}
                .final-table, .param-table {{ width: 100%; border-collapse: collapse; }}
                .final-table tr, .param-table tr {{ border-bottom: 1px solid #e6e6e6; }}
                .final-table td, .param-table td {{ padding: 6px 10px; vertical-align: middle; background: transparent; }}
                .result-label {{ font-weight: bold; color: #333; width: 44%; }}
                .result-value {{ font-weight: bold; color: #2c3e50; }}
                .ok {{ color: #1e8449; font-weight: bold; }}
                .ng {{ color: #c0392b; font-weight: bold; }}
                .formula {{ font-family: 'Consolas', monospace; font-size: 12px; color: #2c3e50; margin: 2px 0; background: transparent; }}
                .formula-result {{ font-family: 'Consolas', monospace; font-size: 12px; color: #1f618d; font-weight: bold; margin: 4px 0 8px 0; padding: 4px 8px; background: #eaf4fb; border-left: 3px solid #2471a3; border-radius: 0 4px 4px 0; }}
                .section-title-bg {{ background: #dceef9; border: 1px solid #aed6f1; border-radius: 6px; padding: 5px 10px; margin: 14px 0 8px 0; font-weight: bold; color: #1f618d; font-size: 13px; }}
                .sub-title {{ border-left: 3px solid #5dade2; padding: 3px 8px; margin: 10px 0 5px 0; font-weight: bold; color: #2471a3; font-size: 12px; background: #f0f8ff; border-radius: 0 4px 4px 0; }}
                .sub-block {{ border-left: 2px solid #aed6f1; margin: 4px 0 8px 12px; padding: 4px 0 4px 10px; background: transparent; }}
                .indent {{ padding-left: 20px; }}
                .check-ok {{ font-family: 'Consolas', monospace; font-size: 12px; font-weight: bold; color: #1e8449; background: #eafaf1; border-left: 3px solid #27ae60; padding: 4px 8px; margin: 4px 0 10px 0; border-radius: 0 4px 4px 0; }}
                .check-ng {{ font-family: 'Consolas', monospace; font-size: 12px; font-weight: bold; color: #c0392b; background: #fdedec; border-left: 3px solid #e74c3c; padding: 4px 8px; margin: 4px 0 10px 0; border-radius: 0 4px 4px 0; }}
                .summary-row {{ font-family: 'Consolas', monospace; font-size: 12px; font-weight: bold; color: #1a5276; background: #d6eaf8; border: 1px solid #aed6f1; border-radius: 4px; padding: 5px 10px; margin: 6px 0 10px 0; }}
            </style>
        </head>
        <body>
            <div class='final-results'>
                <h2>最终计算结果</h2>
                <table class='final-table'>
                    <tr><td class='result-label'>综合结论（按满水+空载控制）</td><td class='result-value {'ok' if c['pass_overall'] else 'ng'}'>{ok(c['pass_overall'])}</td></tr>
                    <tr><td class='result-label'>满水工况安全系数</td><td class='result-value {'ok' if c['pass_full'] else 'ng'}'>{f3(c['K_full'])}</td></tr>
                    <tr><td class='result-label'>实际水深工况安全系数</td><td class='result-value {'ok' if c['pass_actual'] else 'ng'}'>{f3(c['K_actual'])}</td></tr>
                    <tr><td class='result-label'>空载工况安全系数</td><td class='result-value {'ok' if c['pass_empty'] else 'ng'}'>{f3(c['K_empty'])}</td></tr>
                    <tr><td class='result-label'>混凝土总量</td><td class='result-value'>{f3(v['V_concrete_total'])} m³</td></tr>
                    <tr><td class='result-label'>平衡层体积</td><td class='result-value'>{f3(v['V_bal'])} m³</td></tr>
                    <tr><td class='result-label'>换填体积</td><td class='result-value'>{f3(v['V_cushion'])} m³</td></tr>
                    <tr><td class='result-label'>垫层体积</td><td class='result-value'>{f3(v['V_mat'])} m³</td></tr>
                    <tr><td class='result-label'>优化建议</td><td class='result-value'>{suggestion_text}</td></tr>
                </table>
            </div>

            <div class='card'>
                <h2>输入参数</h2>
                <table class='param-table'>
                    <tr><td>内长</td><td>{f3(L)} m</td><td>内宽</td><td>{f3(W)} m</td></tr>
                    <tr><td>池壁总高度</td><td>{f3(H_total)} m</td><td>高出地面高度</td><td>{f3(H_out)} m</td></tr>
                    <tr><td>顶板厚度</td><td>{f3(h_slab)} m</td><td>池壁厚度</td><td>{f3(t_wall)} m</td></tr>
                    <tr><td>底板厚度</td><td>{f3(h_footing)} m</td><td>底板外挑宽度</td><td>{f3(b_overhang)} m</td></tr>
                    <tr><td>有无平衡层</td><td>{'有' if has_bal else '无'}</td><td>平衡层宽出底板距离</td><td>{f3(c_bal_overhang)} m</td></tr>
                    <tr><td>平衡层厚度</td><td>{f3(h_bal)} m</td><td>平衡层材料重度</td><td>{f3(gamma_bal)} kN/m³</td></tr>
                    <tr><td>换填砂石厚度</td><td>{f3(h_cushion)} m</td><td>换填宽出距离</td><td>{f3(d_cushion_overhang)} m</td></tr>
                    <tr><td>换填砂石重度</td><td>{f3(gamma_cushion)} kN/m³</td><td>隔墙总混凝土体积</td><td>{f3(V_partition_total)} m³</td></tr>
                    <tr><td>柱截面长</td><td>{f3(column_a)} m</td><td>柱截面宽</td><td>{f3(column_b)} m</td></tr>
                    <tr><td>柱高类型</td><td>{column_height_type}</td><td>自定义柱高</td><td>{f3(column_height_custom)} m</td></tr>
                    <tr><td>柱个数</td><td>{column_count} 个</td><td>地下水位深度（地面以下）</td><td>{f3(water_table_depth)} m</td></tr>
                    <tr><td>池内实际水深</td><td>{f3(water_depth_inner)} m</td><td>水重度</td><td>{f3(gamma_water)} kN/m³</td></tr>
                    <tr><td>混凝土重度</td><td>{f3(gamma_concrete)} kN/m³</td><td>底板底标高（自动计算）</td><td>{f3(bottom_elev_auto)} m</td></tr>
                    <tr><td>顶板活荷载</td><td>{f3(roof_live_load)} kN/m²</td><td>地面活荷载</td><td>{f3(ground_live_load)} kN/m²</td></tr>
                </table>
            </div>

            <div class='card'>
                <h2>计算过程</h2>

                <div class='section-title-bg'>一、几何与体积计算</div>
                <div class='formula'>内部净高 = 池壁总高度 - 顶板厚度 = {f3(H_total)} - {f3(h_slab)} = {f3(g['H_in'])} m</div>
                <div class='formula'>底板长度 = 内长 + 2×池壁厚度 + 2×底板外挑宽度 = {f3(L)} + 2×{f3(t_wall)} + 2×{f3(b_overhang)} = {f3(g['L_base'])} m</div>
                <div class='formula'>底板宽度 = 内宽 + 2×池壁厚度 + 2×底板外挑宽度 = {f3(W)} + 2×{f3(t_wall)} + 2×{f3(b_overhang)} = {f3(g['W_base'])} m</div>
                <div class='formula'>底板体积 = 底板长度 × 底板宽度 × 底板厚度 = {f3(g['L_base'])} × {f3(g['W_base'])} × {f3(h_footing)} = {f3(v['V_base'])} m³</div>

                <div class='formula'>{bal_formula}</div>
                <div class='formula'>{bal_formula_2}</div>
                <div class='formula'>{bal_volume_formula}</div>

                <div class='formula'>换填长度 = 平衡层长度 + 2×换填宽出距离 = {f3(g['L_bal'])} + 2×{f3(d_cushion_overhang)} = {f3(g['L_cushion'])} m</div>
                <div class='formula'>换填宽度 = 平衡层宽度 + 2×换填宽出距离 = {f3(g['W_bal'])} + 2×{f3(d_cushion_overhang)} = {f3(g['W_cushion'])} m</div>
                <div class='formula'>换填体积 = 换填长度 × 换填宽度 × 换填砂石厚度 = {f3(g['L_cushion'])} × {f3(g['W_cushion'])} × {f3(h_cushion)} = {f3(v['V_cushion'])} m³</div>

                <div class='formula'>池壁中心线总长 = 2×(内长+池壁厚度) + 2×(内宽+池壁厚度) = {f3(g['L_wall_centerline'])} m</div>
                <div class='formula'>池壁总高度 = 水池高度 - 顶板厚度 = {f3(H_total)} - {f3(h_slab)} = {f3(g['H_wall_total'])} m</div>
                <div class='formula'>池壁体积 = 池壁中心线总长 × 池壁厚度 × 池壁总高度 = {f3(g['L_wall_centerline'])} × {f3(t_wall)} × {f3(g['H_wall_total'])} = {f3(v['V_wall'])} m³</div>
                <div class='formula'>顶板长度 = 内长 + 2×池壁厚度 = {f3(L)} + 2×{f3(t_wall)} = {f3(g['L_roof'])} m</div>
                <div class='formula'>顶板宽度 = 内宽 + 2×池壁厚度 = {f3(W)} + 2×{f3(t_wall)} = {f3(g['W_roof'])} m</div>
                <div class='formula'>顶板体积 = 顶板长度 × 顶板宽度 × 顶板厚度 = {f3(g['L_roof'])} × {f3(g['W_roof'])} × {f3(h_slab)} = {f3(v['V_roof'])} m³</div>

                <div class='formula'>{h_column_formula}</div>
                <div class='formula'>单根柱体积 = 柱截面长 × 柱截面宽 × 柱高 = {f3(column_a)} × {f3(column_b)} × {f3(h_column_display)} = {f3(v['V_column_single'])} m³</div>
                <div class='formula'>柱总体积 = 柱个数 × 单根柱体积 = {column_count} × {f3(v['V_column_single'])} = {f3(v['V_columns'])} m³</div>

                <div class='formula'>混凝土总量 = 底板体积 + 池壁体积 + 顶板体积 + 隔墙总混凝土体积 + 柱总体积</div>
                <div class='formula-result'>= {f3(v['V_base'])} + {f3(v['V_wall'])} + {f3(v['V_roof'])} + {f3(v['V_partition_total'])} + {f3(v['V_columns'])} = {f3(v['V_concrete_total'])} m³</div>

                <div class='section-title-bg'>二、抗浮验算过程</div>
                <div class='formula'>底板底埋深 = 池壁总高度 - 高出地面高度 + 底板厚度 = {f3(H_total)} - {f3(H_out)} + {f3(h_footing)} = {f3(bottom_depth)} m</div>
                <div class='formula'>底板底标高 = -底板底埋深 = -{f3(bottom_depth)} = {f3(bottom_elev_auto)} m</div>
                <div class='formula'>地下水位标高 = -地下水位深度（地面以下） = -{f3(water_table_depth)} = {f3(water_table_elev_converted)} m</div>

                <div class='formula'>抗浮受力面积 = {'平衡层长度 × 平衡层宽度' if has_bal else '底板长度 × 底板宽度'} = {f3(g['A_buoy'])} m²</div>
                {f"<div class='formula'>浮力受力底面标高 = 底板底标高 - 平衡层厚度 = {f3(bottom_elev_auto)} - {f3(h_bal)} = {f3(bottom_elev_auto - h_bal)} m（有平衡层，底面取平衡层底）</div>" if has_bal else f"<div class='formula'>浮力受力底面标高 = 底板底标高 = {f3(bottom_elev_auto)} m</div>"}
                <div class='formula'>浮力计算水头 = max(0, 地下水位标高 - 浮力受力底面标高) = max(0, {f3(water_table_elev_converted)} - {f3(bottom_elev_auto - h_bal if has_bal else bottom_elev_auto)}) = {f3(g['H_water_height'])} m</div>
                <div class='formula-result'>浮力 = 抗浮受力面积 × 水重度 × 浮力计算水头 = {f3(w['F_buoy'])} kN</div>

                <div class='formula'>结构自重 = 混凝土总量 × 混凝土重度 = {f3(v['V_concrete_total'])} × {f3(gamma_concrete)} = {f3(w['G_self'])} kN</div>
                <div class='formula'>平衡层重量 = 平衡层体积 × 平衡层材料重度 = {f3(v['V_bal'])} × {f3(gamma_bal)} = {f3(w['G_bal'])} kN</div>
                <div class='formula'>填土有效重度（水位以下）= 饱和重度 - 水重度 = {f3(gamma_soil_sat)} - {f3(gamma_water)} = {f3(w['gamma_soil_eff'])} kN/m³</div>
                <div class='formula'>底板顶标高 = 底板底标高 + 底板厚度 = {f3(bottom_elev_auto)} + {f3(h_footing)} = {f3(bottom_elev_auto + h_footing)} m</div>

                <div class='sub-title'>① 底板外挑环形区填土（从底板顶面起算）</div>
                <div class='sub-block'>
                    <div class='formula'>外挑环形面积 = 底板长×底板宽 - (底板长-2×外挑)×(底板宽-2×外挑)</div>
                    <div class='formula indent'>= {f3(g['L_base'])}×{f3(g['W_base'])} - {f3(g['L_base']-2*b_overhang)}×{f3(g['W_base']-2*b_overhang)} = {f3(w['A_overhang_soil'])} m²</div>
                    <div class='formula'>外挑区填土总高度 = max(0, 0 - {f3(bottom_elev_auto + h_footing)}) = {f3(w['H_overhang_above_wt'] + w['H_overhang_below_wt'])} m</div>
                    <div class='formula indent'>水位以上：{f3(w['H_overhang_above_wt'])} m × 天然重度 {f3(gamma_soil_natural)} kN/m³</div>
                    <div class='formula indent'>水位以下：{f3(w['H_overhang_below_wt'])} m × 有效重度 {f3(w['gamma_soil_eff'])} kN/m³</div>
                    <div class='formula-result'>外挑区填土重量 = {f3(w['A_overhang_soil'])} × ({f3(w['H_overhang_above_wt'])}×{f3(gamma_soil_natural)} + {f3(w['H_overhang_below_wt'])}×{f3(w['gamma_soil_eff'])}) = {f3(w['G_soil_overhang'])} kN</div>
                </div>

                <div class='sub-title'>② 平衡层超出底板环形区填土（从底板底面起算）</div>
                <div class='sub-block'>
                    <div class='formula'>平衡层超出环形面积 = 平衡层长×平衡层宽 - 底板长×底板宽</div>
                    <div class='formula indent'>= {f3(g['L_bal'])}×{f3(g['W_bal'])} - {f3(g['L_base'])}×{f3(g['W_base'])} = {f3(w['A_bal_soil'])} m²</div>
                    <div class='formula'>平衡层区填土总高度 = max(0, 0 - {f3(bottom_elev_auto)}) = {f3(w['H_bal_above_wt'] + w['H_bal_below_wt'])} m</div>
                    <div class='formula indent'>水位以上：{f3(w['H_bal_above_wt'])} m × 天然重度 {f3(gamma_soil_natural)} kN/m³</div>
                    <div class='formula indent'>水位以下：{f3(w['H_bal_below_wt'])} m × 有效重度 {f3(w['gamma_soil_eff'])} kN/m³</div>
                    <div class='formula-result'>平衡层区填土重量 = {f3(w['A_bal_soil'])} × ({f3(w['H_bal_above_wt'])}×{f3(gamma_soil_natural)} + {f3(w['H_bal_below_wt'])}×{f3(w['gamma_soil_eff'])}) = {f3(w['G_soil_bal'])} kN</div>
                </div>

                <div class='formula'>顶板活荷载重量 = 顶板面积 × 顶板活荷载 = {f3(g['L_roof'])} × {f3(g['W_roof'])} × {f3(roof_live_load)} = {f3(w['G_roof_live'])} kN</div>
                <div class='formula'>地面活荷载作用面积 = 底板外挑面积 + {'平衡层外挑面积' if has_bal else '0'} = {f3(w['A_overhang_soil'])} + {f3(w['A_bal_soil'])} = {f3(w['A_ground_live'])} m²</div>
                <div class='formula'>地面活荷载重量 = 地面活荷载作用面积 × 地面活荷载 = {f3(w['A_ground_live'])} × {f3(ground_live_load)} = {f3(w['G_ground_live'])} kN</div>

                <div class='summary-row'>抗浮总重量 = 结构自重 + 平衡层重量 + 外挑区填土 + 平衡层区填土 + 顶板活荷载重量 + 地面活荷载重量<br>&nbsp;&nbsp;= {f3(w['G_self'])} + {f3(w['G_bal'])} + {f3(w['G_soil_overhang'])} + {f3(w['G_soil_bal'])} + {f3(w['G_roof_live'])} + {f3(w['G_ground_live'])} = <b>{f3(w['G_total'])} kN</b></div>

                <div class='formula'>满水重量 = {f3(L)} × {f3(W)} × {f3(g['H_in'])} × {f3(gamma_water)} = {f3(w['G_water_full'])} kN</div>
                <div class='formula'>实际水深重量 = {f3(L)} × {f3(W)} × min({f3(water_depth_inner)}, {f3(g['H_in'])}) × {f3(gamma_water)} = {f3(w['G_water_actual'])} kN</div>

                <div class='formula'>满水工况：K = ({f3(w['G_total'])} + {f3(w['G_water_full'])}) ÷ {f3(w['F_buoy'])} = {k_full_formula}（要求 ≥ {f3(c['req_full'])}）</div>
                <div class='{'check-ok' if c['pass_full'] else 'check-ng'}'>满水工况：{ok(c['pass_full'])}</div>

                <div class='formula'>实际水深工况：K = ({f3(w['G_total'])} + {f3(w['G_water_actual'])}) ÷ {f3(w['F_buoy'])} = {k_actual_formula}（要求 ≥ {f3(c['req_actual'])}）</div>
                <div class='{'check-ok' if c['pass_actual'] else 'check-ng'}'>实际水深工况：{ok(c['pass_actual'])}</div>

                <div class='formula'>空载工况：K = {f3(w['G_total'])} ÷ {f3(w['F_buoy'])} = {k_empty_formula}（要求 ≥ {f3(c['req_empty'])}）</div>
                <div class='{'check-ok' if c['pass_empty'] else 'check-ng'}'>空载工况：{ok(c['pass_empty'])}</div>
            </div>
        </body>
        </html>
        """

    def reset(self):
        """重置输入与输出。"""
        for key, edit in self._inputs.items():
            # 给默认重度保留默认值，其他清空
            if key == "gamma_bal":
                edit.setText("18")
            elif key == "gamma_cushion":
                edit.setText("19")
            elif key == "gamma_water":
                edit.setText("10")
            elif key == "gamma_concrete":
                edit.setText("25")
            elif key in ("gamma_soil_natural", "gamma_soil_sat"):
                edit.setText("18")
            elif key in ("roof_live_load", "ground_live_load"):
                edit.setText("0")
            elif key in ("V_partition_total", "column_a", "column_b", "column_height_custom", "water_depth_inner", "water_table_elev"):
                edit.setText("0")
            else:
                edit.clear()

        self._column_count.setText("0")
        self._has_balancing_layer.setCurrentText("有")
        self._column_height_type.setCurrentText("自定义")
        self._on_balancing_layer_changed()
        self._on_column_height_type_changed()
        self._result.clear()

    def save(self, file_path: str):
        """保存参数到 .fg 文件（JSON）。"""
        data = {k: e.text() for k, e in self._inputs.items()}
        data["has_balancing_layer"] = self._has_balancing_layer.currentText()
        data["column_height_type"] = self._column_height_type.currentText()
        data["column_count"] = self._column_count.text()

        with open(file_path, "w", encoding="utf-8") as f:
            json.dump(data, f, ensure_ascii=False, indent=2)

    def open(self, file_path: str):
        """从 .fg 文件加载参数。"""
        with open(file_path, "r", encoding="utf-8") as f:
            data = json.load(f)

        for k, e in self._inputs.items():
            if k in data:
                e.setText(str(data[k]))

        if "has_balancing_layer" in data:
            self._has_balancing_layer.setCurrentText(str(data["has_balancing_layer"]))
        if "column_height_type" in data:
            self._column_height_type.setCurrentText(str(data["column_height_type"]))
        if "column_count" in data:
            self._column_count.setText(str(data["column_count"]))

        self._on_balancing_layer_changed()
        self._on_column_height_type_changed()

    def _on_export_triggered(self):
        if self._last_result is None:
            QMessageBox.warning(self, "导出错误", "请先进行计算")
            return

        file_path, _ = QFileDialog.getSaveFileName(
            self, "保存计算书", "水池计算书.docx", "Word文档 (*.docx)"
        )
        if not file_path:
            return
        
        # 确保文件路径包含.docx扩展名
        if not file_path.lower().endswith('.docx'):
            file_path += '.docx'

        try:
            doc = Document()
            # 设置中文字体
            doc.styles['Normal'].font.name = '宋体'
            doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            doc.styles['Normal'].font.color.rgb = RGBColor(0, 0, 0)
            doc.styles['Normal'].font.size = Pt(10.5)

            for style_name in ['Title', 'Heading 1', 'Heading 2', 'Heading 3']:
                if style_name in doc.styles:
                    doc.styles[style_name].font.name = '宋体'
                    doc.styles[style_name]._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                    doc.styles[style_name].font.color.rgb = RGBColor(0, 0, 0)

            def f3(value) -> str:
                if isinstance(value, bool):
                    return "是" if value else "否"
                if isinstance(value, int):
                    return str(value)
                if isinstance(value, float):
                    return f"{value:.3f}"
                return str(value)

            input_labels = {
                "L": ("内长", "m"),
                "W": ("内宽", "m"),
                "H_total": ("池壁总高度", "m"),
                "H_out": ("高出地面高度", "m"),
                "h_slab": ("顶板厚度", "m"),
                "t_wall": ("池壁厚度", "m"),
                "h_footing": ("底板厚度", "m"),
                "b_overhang": ("底板外挑宽度", "m"),
                "has_balancing_layer": ("有无平衡层", "-"),
                "c_bal_overhang": ("平衡层宽出底板距离", "m"),
                "h_bal": ("平衡层厚度", "m"),
                "gamma_bal": ("平衡层材料重度", "kN/m³"),
                "h_cushion": ("换填砂石厚度", "m"),
                "d_cushion_overhang": ("换填宽出距离", "m"),
                "gamma_cushion": ("换填砂石重度", "kN/m³"),
                "h_mat": ("垫层厚度", "m"),
                "V_partition_total": ("隔墙总混凝土体积", "m³"),
                "column_a": ("柱截面长", "m"),
                "column_b": ("柱截面宽", "m"),
                "column_height_type": ("柱高类型", "-"),
                "column_height_custom": ("自定义柱高", "m"),
                "column_count": ("柱个数", "个"),
                "water_table_elev": ("地下水位标高", "m"),
                "bottom_elev": ("底板底标高", "m"),
                "gamma_water": ("水重度", "kN/m³"),
                "gamma_concrete": ("混凝土重度", "kN/m³"),
                "water_depth_inner": ("池内实际水深", "m"),
                "gamma_soil_natural": ("填土天然重度", "kN/m³"),
                "gamma_soil_sat": ("填土饱和重度", "kN/m³"),
                "roof_live_load": ("顶板活荷载", "kN/m²"),
                "ground_live_load": ("地面活荷载", "kN/m²"),
            }

            payload = self._last_payload
            result_data = self._last_result
            geometry = result_data.get('geometry', {})
            volumes = result_data.get('volumes', {})
            weights = result_data.get('weights', {})
            checks = result_data.get('checks', {})

            L = payload.L
            W = payload.W
            H_total = payload.H_total
            H_out = payload.H_out
            h_slab = payload.h_slab
            t_wall = payload.t_wall
            h_footing = payload.h_footing
            b_overhang = payload.b_overhang
            c_bal_overhang = payload.c_bal_overhang
            h_bal = payload.h_bal
            gamma_bal = payload.gamma_bal
            h_cushion = payload.h_cushion
            d_cushion_overhang = payload.d_cushion_overhang
            gamma_cushion = payload.gamma_cushion
            h_mat = payload.h_mat
            V_partition_total = payload.V_partition_total
            column_a = payload.column_a
            column_b = payload.column_b
            column_height_custom = payload.column_height_custom
            column_count = payload.column_count
            gamma_water = payload.gamma_water
            gamma_concrete = payload.gamma_concrete
            water_depth_inner = payload.water_depth_inner
            gamma_soil_natural = payload.gamma_soil_natural
            gamma_soil_sat = payload.gamma_soil_sat
            roof_live_load = payload.roof_live_load
            ground_live_load = payload.ground_live_load
            water_table_elev = payload.water_table_elev
            bottom_elev = payload.bottom_elev
            water_table_depth = abs(water_table_elev)
            has_bal = payload.has_balancing_layer
            column_height_type = payload.column_height_type

            def add_text_paragraph(text: str, bold: bool = False):
                p = doc.add_paragraph(style='Normal')
                run = p.add_run(text)
                run.font.name = '宋体'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                run.font.color.rgb = RGBColor(0, 0, 0)
                run.font.size = Pt(10.5)
                run.bold = bold
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.space_before = Pt(0)
                return p

            # 标题
            title = doc.add_heading('水池计算书', level=0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in title.runs:
                run.font.color.rgb = RGBColor(0, 0, 0)
                run.font.name = '宋体'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

            # 日期
            date_str = datetime.now().strftime("%Y-%m-%d")
            date_paragraph = doc.add_paragraph(date_str, style='Normal')
            date_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # 输入参数表格
            doc.add_heading('一、输入参数', level=1)
            input_data = asdict(payload)
            table = doc.add_table(rows=1, cols=2)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = '参数'
            hdr_cells[1].text = '数值'
            for param, value in input_data.items():
                row_cells = table.add_row().cells
                label, unit = input_labels.get(param, (param, '-'))
                row_cells[0].text = label
                if param == 'has_balancing_layer':
                    row_cells[1].text = '有' if value else '无'
                elif param == 'column_height_type':
                    row_cells[1].text = str(value)
                else:
                    row_cells[1].text = f"{f3(value)} {unit}" if unit != '-' else f3(value)

            for row in table.rows:
                row.cells[0].width = Pt(120)
                row.cells[1].width = Pt(220)
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        paragraph.paragraph_format.space_before = Pt(0)
                        paragraph.paragraph_format.space_after = Pt(0)
                        for run in paragraph.runs:
                            run.font.name = '宋体'
                            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                            run.font.color.rgb = RGBColor(0, 0, 0)
                            run.font.size = Pt(10)

            # 计算过程
            doc.add_heading('二、计算过程', level=1)
            doc.add_heading('1. 几何与体积计算', level=2)
            add_text_paragraph(f"(1) 内部净高 H_in = H_total - h_slab = {f3(H_total)} - {f3(h_slab)} = {f3(geometry.get('H_in', 0.0))} m")
            add_text_paragraph(f"(2) 底板长度 L_base = L + 2×t_wall + 2×b_overhang = {f3(L)} + 2×{f3(t_wall)} + 2×{f3(b_overhang)} = {f3(geometry.get('L_base', 0.0))} m")
            add_text_paragraph(f"(3) 底板宽度 W_base = W + 2×t_wall + 2×b_overhang = {f3(W)} + 2×{f3(t_wall)} + 2×{f3(b_overhang)} = {f3(geometry.get('W_base', 0.0))} m")
            add_text_paragraph(f"(4) 底板体积 V_base = L_base × W_base × h_footing = {f3(geometry.get('L_base', 0.0))} × {f3(geometry.get('W_base', 0.0))} × {f3(h_footing)} = {f3(volumes.get('V_base', 0.0))} m³")

            if has_bal:
                add_text_paragraph(f"(5) 平衡层长度 L_bal = L_base + 2×c_bal_overhang = {f3(geometry.get('L_base', 0.0))} + 2×{f3(c_bal_overhang)} = {f3(geometry.get('L_bal', 0.0))} m")
                add_text_paragraph(f"(6) 平衡层宽度 W_bal = W_base + 2×c_bal_overhang = {f3(geometry.get('W_base', 0.0))} + 2×{f3(c_bal_overhang)} = {f3(geometry.get('W_bal', 0.0))} m")
                add_text_paragraph(f"(7) 平衡层体积 V_bal = L_bal × W_bal × h_bal = {f3(geometry.get('L_bal', 0.0))} × {f3(geometry.get('W_bal', 0.0))} × {f3(h_bal)} = {f3(volumes.get('V_bal', 0.0))} m³")
            else:
                add_text_paragraph(f"(5) 本工程无平衡层，因此 L_bal = L_base = {f3(geometry.get('L_bal', 0.0))} m，W_bal = W_base = {f3(geometry.get('W_bal', 0.0))} m")
                add_text_paragraph("(6) 本工程无平衡层，因此平衡层体积 V_bal = 0.000 m³")

            add_text_paragraph(f"(8) 换填长度 L_cushion = {f3(geometry.get('L_cushion', 0.0))} m，换填宽度 W_cushion = {f3(geometry.get('W_cushion', 0.0))} m")
            add_text_paragraph(f"(9) 换填体积 V_cushion = L_cushion × W_cushion × h_cushion = {f3(geometry.get('L_cushion', 0.0))} × {f3(geometry.get('W_cushion', 0.0))} × {f3(h_cushion)} = {f3(volumes.get('V_cushion', 0.0))} m³")
            add_text_paragraph(f"(10) 垫层体积 V_mat = (L_base+0.2) × (W_base+0.2) × h_mat，对应计算结果为 {f3(volumes.get('V_mat', 0.0))} m³")
            add_text_paragraph(f"(11) 池壁中心线总长 = 2×(L+t_wall)+2×(W+t_wall) = {f3(geometry.get('L_wall_centerline', 0.0))} m")
            add_text_paragraph(f"(12) 池壁总高度 H_wall_total = H_total - h_slab = {f3(H_total)} - {f3(h_slab)} = {f3(geometry.get('H_wall_total', 0.0))} m")
            add_text_paragraph(f"(13) 池壁体积 V_wall = 池壁中心线总长 × t_wall × H_wall_total = {f3(geometry.get('L_wall_centerline', 0.0))} × {f3(t_wall)} × {f3(geometry.get('H_wall_total', 0.0))} = {f3(volumes.get('V_wall', 0.0))} m³")
            add_text_paragraph(f"(14) 顶板长度 L_roof = L + 2×t_wall = {f3(L)} + 2×{f3(t_wall)} = {f3(geometry.get('L_roof', 0.0))} m")
            add_text_paragraph(f"(15) 顶板宽度 W_roof = W + 2×t_wall = {f3(W)} + 2×{f3(t_wall)} = {f3(geometry.get('W_roof', 0.0))} m")
            add_text_paragraph(f"(16) 顶板体积 V_roof = L_roof × W_roof × h_slab = {f3(geometry.get('L_roof', 0.0))} × {f3(geometry.get('W_roof', 0.0))} × {f3(h_slab)} = {f3(volumes.get('V_roof', 0.0))} m³")
            if column_height_type == '内净高':
                add_text_paragraph(f"(17) 柱高取内部净高，即 H_column = {f3(geometry.get('H_column', 0.0))} m")
            else:
                add_text_paragraph(f"(17) 柱高取自定义值，即 H_column = {f3(column_height_custom)} m")
            add_text_paragraph(f"(18) 单根柱体积 V_column_single = column_a × column_b × H_column = {f3(column_a)} × {f3(column_b)} × {f3(geometry.get('H_column', 0.0))} = {f3(volumes.get('V_column_single', 0.0))} m³")
            add_text_paragraph(f"(19) 柱总体积 V_columns = column_count × V_column_single = {column_count} × {f3(volumes.get('V_column_single', 0.0))} = {f3(volumes.get('V_columns', 0.0))} m³")
            add_text_paragraph(f"(20) 混凝土总量 V_concrete_total = V_base + V_wall + V_roof + V_partition_total + V_columns = {f3(volumes.get('V_base', 0.0))} + {f3(volumes.get('V_wall', 0.0))} + {f3(volumes.get('V_roof', 0.0))} + {f3(V_partition_total)} + {f3(volumes.get('V_columns', 0.0))} = {f3(volumes.get('V_concrete_total', 0.0))} m³")

            doc.add_heading('2. 抗浮验算过程', level=2)
            add_text_paragraph(f"(1) 底板底埋深 = H_total - H_out + h_footing = {f3(H_total)} - {f3(H_out)} + {f3(h_footing)} = {f3(abs(bottom_elev))} m")
            add_text_paragraph(f"(2) 底板底标高 bottom_elev = -底板底埋深 = {f3(bottom_elev)} m")
            add_text_paragraph(f"(3) 地下水位深度输入值为 {f3(water_table_depth)} m（地面以下），换算地下水位标高 water_table_elev = -{f3(water_table_depth)} = {f3(water_table_elev)} m")
            add_text_paragraph(f"(4) 抗浮受力面积 A_buoy = {'L_bal × W_bal' if has_bal else 'L_base × W_base'} = {f3(geometry.get('A_buoy', 0.0))} m²")
            if has_bal:
                add_text_paragraph(f"(5) 有平衡层时，浮力受力底面标高 = bottom_elev - h_bal = {f3(bottom_elev)} - {f3(h_bal)} = {f3(bottom_elev - h_bal)} m")
            else:
                add_text_paragraph(f"(5) 无平衡层时，浮力受力底面标高 = bottom_elev = {f3(bottom_elev)} m")
            add_text_paragraph(f"(6) 浮力计算水头 H_water_height = max(0, water_table_elev - 浮力受力底面标高) = {f3(geometry.get('H_water_height', 0.0))} m")
            add_text_paragraph(f"(7) 浮力 F_buoy = A_buoy × gamma_water × H_water_height = {f3(geometry.get('A_buoy', 0.0))} × {f3(gamma_water)} × {f3(geometry.get('H_water_height', 0.0))} = {f3(weights.get('F_buoy', 0.0))} kN")
            add_text_paragraph(f"(8) 结构自重 G_self = V_concrete_total × gamma_concrete = {f3(volumes.get('V_concrete_total', 0.0))} × {f3(gamma_concrete)} = {f3(weights.get('G_self', 0.0))} kN")
            add_text_paragraph(f"(9) 平衡层重量 G_bal = V_bal × gamma_bal = {f3(volumes.get('V_bal', 0.0))} × {f3(gamma_bal)} = {f3(weights.get('G_bal', 0.0))} kN")
            add_text_paragraph(f"(11) 填土有效重度 gamma_soil_eff = gamma_soil_sat - gamma_water = {f3(gamma_soil_sat)} - {f3(gamma_water)} = {f3(weights.get('gamma_soil_eff', 0.0))} kN/m³")
            add_text_paragraph(f"(12) 底板外挑环形区面积 A_overhang_soil = {f3(weights.get('A_overhang_soil', 0.0))} m²")
            add_text_paragraph(f"(13) 外挑区水位以上填土高度 = {f3(weights.get('H_overhang_above_wt', 0.0))} m，水位以下填土高度 = {f3(weights.get('H_overhang_below_wt', 0.0))} m")
            add_text_paragraph(f"(14) 外挑区填土重量 G_soil_overhang = A_overhang_soil × (H_above×gamma_soil_natural + H_below×gamma_soil_eff) = {f3(weights.get('G_soil_overhang', 0.0))} kN")
            add_text_paragraph(f"(15) 平衡层超出底板环形区面积 A_bal_soil = {f3(weights.get('A_bal_soil', 0.0))} m²")
            add_text_paragraph(f"(16) 平衡层区水位以上填土高度 = {f3(weights.get('H_bal_above_wt', 0.0))} m，水位以下填土高度 = {f3(weights.get('H_bal_below_wt', 0.0))} m")
            add_text_paragraph(f"(17) 平衡层区填土重量 G_soil_bal = {f3(weights.get('G_soil_bal', 0.0))} kN")
            add_text_paragraph(f"(18) 顶板活荷载重量 G_roof_live = L_roof × W_roof × roof_live_load = {f3(geometry.get('L_roof', 0.0))} × {f3(geometry.get('W_roof', 0.0))} × {f3(roof_live_load)} = {f3(weights.get('G_roof_live', 0.0))} kN")
            add_text_paragraph(f"(19) 地面活荷载作用面积 A_ground_live = A_overhang_soil + {'A_bal_soil' if has_bal else '0'} = {f3(weights.get('A_overhang_soil', 0.0))} + {f3(weights.get('A_bal_soil', 0.0))} = {f3(weights.get('A_ground_live', 0.0))} m²")
            add_text_paragraph(f"(20) 地面活荷载重量 G_ground_live = A_ground_live × ground_live_load = {f3(weights.get('A_ground_live', 0.0))} × {f3(ground_live_load)} = {f3(weights.get('G_ground_live', 0.0))} kN")
            add_text_paragraph(f"(21) 抗浮总重量 G_total = G_self + G_bal + G_soil_overhang + G_soil_bal + G_roof_live + G_ground_live = {f3(weights.get('G_self', 0.0))} + {f3(weights.get('G_bal', 0.0))} + {f3(weights.get('G_soil_overhang', 0.0))} + {f3(weights.get('G_soil_bal', 0.0))} + {f3(weights.get('G_roof_live', 0.0))} + {f3(weights.get('G_ground_live', 0.0))} = {f3(weights.get('G_total', 0.0))} kN")
            add_text_paragraph(f"(19) 满水重量 G_water_full = L × W × H_in × gamma_water = {f3(L)} × {f3(W)} × {f3(geometry.get('H_in', 0.0))} × {f3(gamma_water)} = {f3(weights.get('G_water_full', 0.0))} kN")
            add_text_paragraph(f"(20) 实际水深重量 G_water_actual = L × W × h_water_actual × gamma_water = {f3(L)} × {f3(W)} × {f3(geometry.get('h_water_actual', 0.0))} × {f3(gamma_water)} = {f3(weights.get('G_water_actual', 0.0))} kN")
            add_text_paragraph(f"(21) 满水工况安全系数 K_full = (G_total + G_water_full) ÷ F_buoy = ({f3(weights.get('G_total', 0.0))} + {f3(weights.get('G_water_full', 0.0))}) ÷ {f3(weights.get('F_buoy', 0.0))} = {f3(checks.get('K_full', 0.0))}，要求 ≥ {f3(checks.get('req_full', 0.0))}")
            add_text_paragraph(f"(22) 实际水深工况安全系数 K_actual = (G_total + G_water_actual) ÷ F_buoy = ({f3(weights.get('G_total', 0.0))} + {f3(weights.get('G_water_actual', 0.0))}) ÷ {f3(weights.get('F_buoy', 0.0))} = {f3(checks.get('K_actual', 0.0))}，要求 ≥ {f3(checks.get('req_actual', 0.0))}")
            add_text_paragraph(f"(23) 空载工况安全系数 K_empty = G_total ÷ F_buoy = {f3(weights.get('G_total', 0.0))} ÷ {f3(weights.get('F_buoy', 0.0))} = {f3(checks.get('K_empty', 0.0))}，要求 ≥ {f3(checks.get('req_empty', 0.0))}")

            # 最终结果表格
            doc.add_heading('三、最终计算结果', level=1)
            result_table = doc.add_table(rows=1, cols=2)
            result_table.style = 'Table Grid'
            hdr_cells = result_table.rows[0].cells
            hdr_cells[0].text = '项目'
            hdr_cells[1].text = '结果'

            summary_rows = [
                ('综合结论（按满水+空载控制）', '满足' if checks.get('pass_overall') else '不满足'),
                ('满水工况安全系数', f3(checks.get('K_full', 0.0))),
                ('实际水深工况安全系数', f3(checks.get('K_actual', 0.0))),
                ('空载工况安全系数', f3(checks.get('K_empty', 0.0))),
                ('混凝土总量', f"{f3(volumes.get('V_concrete_total', 0.0))} m³"),
                ('平衡层体积', f"{f3(volumes.get('V_bal', 0.0))} m³"),
                ('换填体积', f"{f3(volumes.get('V_cushion', 0.0))} m³"),
                ('垫层体积', f"{f3(volumes.get('V_mat', 0.0))} m³"),
            ]

            for item, value in summary_rows:
                row_cells = result_table.add_row().cells
                row_cells[0].text = item
                row_cells[1].text = value

            for row in result_table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        paragraph.paragraph_format.space_before = Pt(0)
                        paragraph.paragraph_format.space_after = Pt(0)
                        for run in paragraph.runs:
                            run.font.name = '宋体'
                            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                            run.font.color.rgb = RGBColor(0, 0, 0)
                            run.font.size = Pt(10)

            doc.save(file_path)
        except Exception as e:
            QMessageBox.critical(self, "保存失败", f"文档生成失败: {str(e)}")
            return

        # 导出成功提示
        msg_box = QMessageBox(self)
        msg_box.setWindowTitle("导出成功")
        msg_box.setText(f"计算书已保存至:\n{file_path}")
        msg_box.addButton("打开文件", QMessageBox.AcceptRole)
        msg_box.addButton("确定", QMessageBox.RejectRole)
        ret = msg_box.exec_()
        if ret == 0:
            QDesktopServices.openUrl(QUrl.fromLocalFile(file_path))
